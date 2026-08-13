#!/usr/bin/env python3
"""Build a Word deliverable through a Markdown/YAML/BibTeX -> DOCX pipeline.

Sibling of ``build_latex_report.py``: same entry shape (a folder holding
``report.yml``), same configuration object, same global-publication step. The
block grammar it accepts is deliberately the one ``markdown_to_latex()``
defines, so a single ``body.md`` renders through either backend.

Three limits are structural to Word and are handled explicitly rather than
silently:

* **Equations.** ``python-docx`` writes no OMML, so ``$...$`` and ``$$...$$``
  are rendered as their LaTeX SOURCE in a distinct monospace style and the
  build warns with a count. Mangling them into prose is what
  ``common_validation`` already treats as a critical defect; refusing the whole
  build over one inline ``$x$`` would be disproportionate for a format the
  skill offers as a delivery option.
* **Raw LaTeX passthrough.** ``\\begin{...}``/``\\end{...}`` lines mean nothing
  here. They are skipped with a named warning instead of printed literally.
* **Font fallbacks.** OOXML has no fallback list; substitution happens through
  ``w:altName`` in the font table, which is what ``declare_font_fallbacks``
  writes from ``academic_format.yml``.
"""
from __future__ import annotations

import argparse
import os
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from lxml import etree

from build_latex_report import (
    ASSETS_DIR,
    FIGURE_SUFFIXES,
    LOGO_FILENAME,
    is_bibliography_heading,
)
from output_router import publish_global_output
from report_config import ReportConfig, load_report_config

# ---------------------------------------------------------------------------
# Style names and page geometry
# ---------------------------------------------------------------------------

MONO_FONT = "Courier New"
CODE_STYLE = "Código fuente"
MATH_STYLE = "Fórmula"
CAPTION_STYLE = "Caption"
REFERENCE_STYLE = "Referencia"
COVER_STYLE = "Portada"

# Page sizes the builder knows how to set, in centimetres.
PAGE_SIZES: dict[str, tuple[float, float]] = {
    "a4": (21.0, 29.7),
    "letter": (21.59, 27.94),
}

HEADING_SIZES_PT = {1: 16, 2: 14, 3: 13, 4: 12}

# Institutional strings. Route A only — see references/document-routing.md.
DEFAULT_UNIVERSITY = "Universidad Nacional de Loja"
DEFAULT_FACULTY = (
    "Facultad de la Energía, las Industrias y los Recursos Naturales no Renovables"
)
COVER_TABLE_FIELDS = (
    ("subject", "Asignatura"),
    ("teacher", "Docente"),
    ("student", "Estudiante"),
    ("parallel", "Paralelo"),
    ("career", "Carrera"),
    ("date", "Fecha"),
)

DEFAULT_BIBLIOGRAPHY_TITLE = "Referencias"


# ---------------------------------------------------------------------------
# Inline grammar
# ---------------------------------------------------------------------------
#
# Same precedence as convert_inline() in build_latex_report: citations, math and
# inline code are recognised before the emphasis markers, so a `**` inside a
# code span or a formula is never mistaken for bold.
INLINE_RE = re.compile(
    r"\[@(?P<cite>[A-Za-z0-9_:\-.,; ]+)\]"
    r"|\$(?P<math>[^$]+)\$"
    r"|`(?P<code>[^`]+)`"
    r"|\*\*(?P<bold>[^*]+)\*\*"
    r"|(?<!\*)\*(?P<italic>[^*]+)\*(?!\*)"
)

CITATION_RE = re.compile(r"\[@([A-Za-z0-9_:\-.,; ]+)\]")
CODE_FENCE_RE = re.compile(r"^(?P<marker>`{3,}|~{3,})\s*(?P<language>[^`]*)$")
HEADING_RE = re.compile(r"^(#{1,4})\s+(.+)$")
IMAGE_RE = re.compile(r"^!\[(?P<caption>[^\]]*)\]\((?P<src>[^)]+)\)")
BULLET_RE = re.compile(r"^[-*]\s+(.+)$")
ORDERED_RE = re.compile(r"^\d+[.)]\s+(.+)$")
DISPLAY_MATH_RE = re.compile(r"^\$\$(?P<formula>.+)\$\$$")
PAGE_BREAK_COMMANDS = {"\\newpage", "\\clearpage", "\\pagebreak"}
LATEX_ENVIRONMENT_RE = re.compile(r"^\\(begin|end)\{")


@dataclass(frozen=True)
class Segment:
    """One inline run: its text plus the formatting it carries."""

    text: str
    bold: bool = False
    italic: bool = False
    mono: bool = False


def split_citation_keys(raw: str) -> list[str]:
    """Split the inside of ``[@a, b; c]`` into its keys, in written order."""
    return [key.strip() for key in re.split(r"[,;]", raw) if key.strip()]


def cited_keys_in_order(markdown: str) -> list[str]:
    """Every cited key, deduplicated, in order of first appearance."""
    ordered: list[str] = []
    for match in CITATION_RE.finditer(markdown):
        for key in split_citation_keys(match.group(1)):
            if key not in ordered:
                ordered.append(key)
    return ordered


def _merge_flags(
    segments: Iterable[Segment], *, bold: bool = False, italic: bool = False
) -> list[Segment]:
    """Re-emit segments carrying an outer flag merged onto their own flags.

    The emphasis group in ``INLINE_RE`` cannot itself contain ``*``, so the
    recursion inside a bold/italic span only ever meets code, math or citation
    rules — nesting stays at depth one by construction.
    """
    return [
        Segment(
            segment.text,
            bold=bold or segment.bold,
            italic=italic or segment.italic,
            mono=segment.mono,
        )
        for segment in segments
    ]


def inline_segments(text: str, citation_numbers: dict[str, int] | None = None) -> list[Segment]:
    """Split a line of Markdown into formatted runs.

    ``citation_numbers`` maps a BibTeX key to its IEEE number. A key missing
    from that map cannot happen — ``resolve_citations`` refuses the build first
    — so there is no branch here that emits a marker pointing at nothing.

    Inner rules (code, math, cite) are re-parsed inside a bold or italic span
    so the segment stream carries every applicable flag instead of literal
    Markdown markers; the outer emphasis flag merges onto each inner segment.
    """
    numbers = citation_numbers or {}
    segments: list[Segment] = []
    cursor = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > cursor:
            segments.append(Segment(text[cursor : match.start()]))
        cursor = match.end()
        if match.group("cite") is not None:
            keys = split_citation_keys(match.group("cite"))
            segments.append(Segment(", ".join(f"[{numbers[key]}]" for key in keys)))
        elif match.group("math") is not None:
            segments.append(Segment(match.group("math"), mono=True))
        elif match.group("code") is not None:
            segments.append(Segment(match.group("code"), mono=True))
        elif match.group("bold") is not None:
            segments.extend(
                _merge_flags(inline_segments(match.group("bold"), numbers), bold=True)
            )
        else:
            segments.extend(
                _merge_flags(inline_segments(match.group("italic"), numbers), italic=True)
            )
    if cursor < len(text):
        segments.append(Segment(text[cursor:]))
    return [segment for segment in segments if segment.text]


def count_math_spans(markdown: str) -> int:
    """How many ``$...$`` / ``$$...$$`` spans the body carries."""
    return len(re.findall(r"\$\$.+?\$\$|\$[^$]+\$", markdown, re.S))


# ---------------------------------------------------------------------------
# BibTeX
# ---------------------------------------------------------------------------


def _strip_latex(value: str) -> str:
    """Reduce a BibTeX field to plain text Word can show."""
    text = re.sub(r"\\url\{([^}]*)\}", r"\1", value)
    text = re.sub(r"\\[a-zA-Z]+\s*", "", text)
    text = text.replace("{", "").replace("}", "")
    return " ".join(text.split())


def parse_bib(bib_text: str) -> dict[str, dict[str, str]]:
    """Parse just enough BibTeX to render an IEEE-style reference list.

    Deliberately minimal: entry type, key and top-level ``field = {...}`` pairs
    with balanced braces. Anything richer belongs to biber, which the LaTeX
    backend already uses.
    """
    entries: dict[str, dict[str, str]] = {}
    for match in re.finditer(r"@(\w+)\s*\{\s*([^,\s]+)\s*,", bib_text):
        key = match.group(2)
        fields: dict[str, str] = {"entrytype": match.group(1).lower()}
        cursor = match.end()
        depth = 0
        # Walk the entry body, tracking brace depth so a nested {} never ends it.
        while cursor < len(bib_text):
            character = bib_text[cursor]
            if character == "{":
                depth += 1
            elif character == "}":
                if depth == 0:
                    break
                depth -= 1
            cursor += 1
        body = bib_text[match.end() : cursor]
        for field in re.finditer(
            r"(\w+)\s*=\s*(\{(?:[^{}]|\{[^{}]*\})*\}|\"[^\"]*\"|[^,\n]+)", body
        ):
            fields[field.group(1).strip().lower()] = _strip_latex(
                field.group(2).strip().strip("{}\"")
            )
        entries[key] = fields
    return entries


def format_authors(raw: str) -> str:
    """Render ``Last, First and Other, Second`` as ``F. Last, S. Other``."""
    people: list[str] = []
    for author in re.split(r"\s+and\s+", raw):
        author = author.strip()
        if not author:
            continue
        if "," in author:
            last, _, first = author.partition(",")
        else:
            parts = author.split()
            last, first = parts[-1], " ".join(parts[:-1])
        initials = " ".join(f"{part[0]}." for part in first.split() if part)
        people.append(f"{initials} {last.strip()}".strip())
    return ", ".join(people)


def format_reference(number: int, key: str, fields: dict[str, str]) -> str:
    """One IEEE-shaped reference line."""
    parts: list[str] = [f"[{number}]"]
    author = format_authors(fields.get("author", "")) or fields.get("organization", "")
    if author:
        parts.append(f"{author},")
    title = fields.get("title", key)
    parts.append(f'"{title},"')
    container = (
        fields.get("journal")
        or fields.get("booktitle")
        or fields.get("publisher")
        or fields.get("howpublished")
        or ""
    )
    if container:
        parts.append(f"{container},")
    year = fields.get("year", "")
    parts.append(f"{year}." if year else "")
    tail = fields.get("doi") or fields.get("url") or ""
    if tail:
        parts.append(f"doi: {tail}." if fields.get("doi") else tail)
    return " ".join(part for part in parts if part)


def resolve_citations(
    markdown: str, bib_text: str, bib_path: Path | None
) -> tuple[dict[str, int], list[tuple[int, str, dict[str, str]]], list[str]]:
    """Number the cited works and report both failure modes honestly.

    Mirrors the severities ``validate_ieee_refs`` already encodes for this
    project: a citation with no BibTeX entry (or no bibliography at all) is a
    hard error, an entry that nobody cites is a warning.
    """
    cited = cited_keys_in_order(markdown)
    if not cited:
        entries = parse_bib(bib_text)
        unused = sorted(entries) if entries else []
        return {}, [], unused

    if bib_path is None:
        raise SystemExit(
            "Hay citas en el cuerpo, pero no existe sources.bib/bibliography "
            "configurada. Claves citadas: " + ", ".join(cited)
        )

    entries = parse_bib(bib_text)
    missing = [key for key in cited if key not in entries]
    if missing:
        raise SystemExit(
            f"Citas sin entrada BibTeX en {bib_path}: " + ", ".join(sorted(missing))
        )

    numbers = {key: index for index, key in enumerate(cited, start=1)}
    bibliography = [(numbers[key], key, entries[key]) for key in cited]
    unused = sorted(key for key in entries if key not in numbers)
    return numbers, bibliography, unused


# ---------------------------------------------------------------------------
# Figures
# ---------------------------------------------------------------------------


def figure_search_dirs(config: ReportConfig) -> list[Path]:
    r"""Directories a figure reference may resolve against, in priority order.

    The build directory comes first because that is what ``\includegraphics``
    uses, and it is why real bodies carry ``../../../assets/generated/...``.
    The report folder and the Markdown file's own directory follow, so a body
    written for DOCX only can also use the path a reader would expect.
    """
    candidates = [config.tex_path.parent, config.folder, config.body_path.parent]
    unique: list[Path] = []
    for candidate in candidates:
        if candidate not in unique:
            unique.append(candidate)
    return unique


def resolve_figure(reference: str, search_dirs: Iterable[Path]) -> Path | None:
    """Resolve a figure reference, trying the graphics suffixes when absent.

    The joined path is normalised lexically because a report often has no
    ``build/`` directory yet, and the kernel refuses to walk ``..`` out of a
    directory that does not exist — which would report every
    ``../../../assets/...`` reference as missing.
    """
    candidate = Path(reference)
    bases = (
        [candidate]
        if candidate.is_absolute()
        else [Path(os.path.normpath(directory / candidate)) for directory in search_dirs]
    )
    for base in bases:
        if base.suffix:
            if base.exists():
                return base
            continue
        for suffix in FIGURE_SUFFIXES:
            with_suffix = base.with_suffix(suffix)
            if with_suffix.exists():
                return with_suffix
    return None


# ---------------------------------------------------------------------------
# Document scaffolding
# ---------------------------------------------------------------------------


def page_size_cm(config: ReportConfig) -> tuple[float, float]:
    name = str(config.academic_value("page", "size", default="A4")).strip().lower()
    if name not in PAGE_SIZES:
        raise SystemExit(
            f"Tamaño de página no soportado por el builder DOCX: {name!r}. "
            "Valores válidos: " + ", ".join(sorted(PAGE_SIZES))
        )
    return PAGE_SIZES[name]


def apply_page_setup(document: Document, config: ReportConfig) -> None:
    width_cm, height_cm = page_size_cm(config)
    margin_cm = float(config.academic_value("page", "margins", "default_cm", default=2.5))
    for section in document.sections:
        section.page_width = Cm(width_cm)
        section.page_height = Cm(height_cm)
        section.top_margin = Cm(margin_cm)
        section.bottom_margin = Cm(margin_cm)
        section.left_margin = Cm(margin_cm)
        section.right_margin = Cm(margin_cm)


def font_preferences(config: ReportConfig) -> tuple[str, str]:
    """The serif body font and the sans font for figure labels."""
    preference = config.academic_value("body_text", "font_preference", default=[]) or []
    serif = str(preference[0]) if len(preference) > 0 else "Times New Roman"
    sans = str(preference[1]) if len(preference) > 1 else "Arial"
    return serif, sans


def set_run_font(run, name: str, size_pt: float | None = None) -> None:
    """Bind a run to one font family across every script slot."""
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attribute in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attribute), name)
    if size_pt is not None:
        run.font.size = Pt(size_pt)


def declare_font_fallbacks(document: Document, config: ReportConfig) -> None:
    """Write the documented fallback chain into ``word/fontTable.xml``.

    OOXML carries no CSS-style font list: a run names exactly one family, and
    the renderer substitutes through ``w:altName`` when that family is absent.
    Writing the chain there is therefore the only way a Linux LibreOffice shows
    TeX Gyre Termes instead of an arbitrary default.

    A rendering hint must never break a build, so an unexpected package layout
    is a silent no-op.
    """
    serif, sans = font_preferences(config)
    chain = config.academic_value("body_text", "latex_fallback_chain", default={}) or {}
    alternatives = [
        (serif, (chain.get("serif") or ["TeX Gyre Termes"])[0]),
        (sans, (chain.get("sans") or ["TeX Gyre Heros"])[0]),
    ]
    linux_fallbacks = config.academic_value("body_text", "linux_pdf_fallbacks", default={}) or {}

    part = None
    for candidate in document.part.package.iter_parts():
        if str(candidate.partname) == "/word/fontTable.xml":
            part = candidate
            break
    if part is None or not hasattr(part, "_blob"):
        return
    try:
        root = etree.fromstring(part.blob)
    except etree.XMLSyntaxError:
        return
    for family, alternative in alternatives:
        font = etree.SubElement(root, qn("w:font"))
        font.set(qn("w:name"), family)
        alt = etree.SubElement(font, qn("w:altName"))
        alt.set(qn("w:val"), alternative)
    for family, alternative in (
        (serif, linux_fallbacks.get("serif")),
        (sans, linux_fallbacks.get("sans")),
    ):
        if not alternative:
            continue
        font = etree.SubElement(root, qn("w:font"))
        font.set(qn("w:name"), family)
        alt = etree.SubElement(font, qn("w:altName"))
        alt.set(qn("w:val"), str(alternative))
    part._blob = etree.tostring(
        root, xml_declaration=True, encoding="UTF-8", standalone=True
    )


def ensure_paragraph_style(document: Document, name: str):
    from docx.enum.style import WD_STYLE_TYPE

    try:
        return document.styles[name]
    except KeyError:
        return document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)


def apply_styles(document: Document, config: ReportConfig) -> None:
    """Bind every style this builder uses to the format contract."""
    serif, sans = font_preferences(config)
    size_pt = float(config.academic_value("body_text", "size_pt", default=12))
    line_spacing = float(config.academic_value("body_text", "line_spacing", default=1.15))
    indent_cm = float(
        config.academic_value("body_text", "first_line_indent_cm", default=1.27)
    )
    alignment = str(
        config.academic_value("body_text", "alignment_default", default="justified")
    ).strip().lower()

    normal = document.styles["Normal"]
    normal.font.name = serif
    normal.font.size = Pt(size_pt)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), serif)
    normal.element.rPr.rFonts.set(qn("w:cs"), serif)
    fmt = normal.paragraph_format
    fmt.alignment = (
        WD_ALIGN_PARAGRAPH.JUSTIFY if alignment == "justified" else WD_ALIGN_PARAGRAPH.LEFT
    )
    fmt.line_spacing = line_spacing
    fmt.first_line_indent = Cm(indent_cm)
    fmt.space_after = Pt(6)

    # Everything that is not running prose must not inherit the paragraph indent.
    for name in ("List Bullet", "List Number", CAPTION_STYLE, "Title"):
        style = ensure_paragraph_style(document, name)
        style.paragraph_format.first_line_indent = Cm(0)

    for level, heading_size in HEADING_SIZES_PT.items():
        style = document.styles[f"Heading {level}"]
        style.font.name = serif
        style.font.size = Pt(heading_size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        style.paragraph_format.first_line_indent = Cm(0)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        style.paragraph_format.space_before = Pt(12 if level <= 2 else 8)
        style.paragraph_format.space_after = Pt(6)

    caption = document.styles[CAPTION_STYLE]
    caption.font.name = sans
    caption.font.size = Pt(10)
    caption.font.italic = False
    caption.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(4)

    code = ensure_paragraph_style(document, CODE_STYLE)
    code.font.name = MONO_FONT
    code.font.size = Pt(9.5)
    code.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    code.paragraph_format.first_line_indent = Cm(0)
    code.paragraph_format.left_indent = Cm(0.5)
    code.paragraph_format.line_spacing = 1.0
    code.paragraph_format.space_before = Pt(6)
    code.paragraph_format.space_after = Pt(6)

    math = ensure_paragraph_style(document, MATH_STYLE)
    math.font.name = MONO_FONT
    math.font.size = Pt(11)
    math.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    math.paragraph_format.first_line_indent = Cm(0)
    math.paragraph_format.line_spacing = 1.0

    reference = ensure_paragraph_style(document, REFERENCE_STYLE)
    reference.font.name = serif
    reference.font.size = Pt(11)
    reference.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    reference.paragraph_format.first_line_indent = Cm(-0.75)
    reference.paragraph_format.left_indent = Cm(0.75)
    reference.paragraph_format.space_after = Pt(6)

    cover = ensure_paragraph_style(document, COVER_STYLE)
    cover.font.name = serif
    cover.font.size = Pt(12)
    cover.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover.paragraph_format.first_line_indent = Cm(0)
    cover.paragraph_format.line_spacing = 1.0
    cover.paragraph_format.space_after = Pt(4)


def set_paragraph_shading(paragraph, fill: str) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    ppr.append(shd)


def paragraph_num_id(paragraph) -> int | None:
    """The explicit ``w:numId`` bound to a paragraph, if any."""
    ppr = paragraph._p.pPr
    if ppr is None:
        return None
    num_pr = ppr.find(qn("w:numPr"))
    if num_pr is None:
        return None
    num_id = num_pr.find(qn("w:numId"))
    if num_id is None:
        return None
    value = num_id.get(qn("w:val"))
    return int(value) if value is not None else None


def _style_num_id(document: Document, style_name: str) -> int | None:
    style = document.styles[style_name]
    ppr = style.element.find(qn("w:pPr"))
    if ppr is None:
        return None
    num_pr = ppr.find(qn("w:numPr"))
    if num_pr is None:
        return None
    num_id = num_pr.find(qn("w:numId"))
    if num_id is None:
        return None
    value = num_id.get(qn("w:val"))
    return int(value) if value is not None else None


def new_numbering_id(document: Document, style_name: str = "List Number") -> int | None:
    """Clone the list style's numbering so the next list restarts at 1.

    Word continues a numbering definition wherever it is reused, so two
    procedures in one document would run 1..2 then 3..4. Each ordered list gets
    its own ``w:num`` pointing at the same abstract definition, and an explicit
    ``w:lvlOverride``/``w:startOverride val="1"`` forces that instance to start
    counting again — the restart mechanism OOXML actually offers. The override
    must follow ``w:abstractNumId`` to satisfy the ``CT_Num`` sequence.
    """
    numbering = document.part.numbering_part.element
    base_num_id = _style_num_id(document, style_name)
    if base_num_id is None:
        return None
    abstract_id = None
    used: list[int] = []
    for num in numbering.findall(qn("w:num")):
        value = num.get(qn("w:numId"))
        if value is None:
            continue
        used.append(int(value))
        if int(value) == base_num_id:
            abstract = num.find(qn("w:abstractNumId"))
            if abstract is not None:
                abstract_id = abstract.get(qn("w:val"))
    if abstract_id is None:
        return None
    next_id = max(used) + 1 if used else 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(next_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), abstract_id)
    num.append(abstract)
    level_override = OxmlElement("w:lvlOverride")
    level_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    level_override.append(start_override)
    num.append(level_override)
    numbering.append(num)
    return next_id


def bind_numbering(paragraph, num_id: int) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_pr.append(ilvl)
    num_element = OxmlElement("w:numId")
    num_element.set(qn("w:val"), str(num_id))
    num_pr.append(num_element)
    ppr.append(num_pr)


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


# ---------------------------------------------------------------------------
# Renderer
# ---------------------------------------------------------------------------


class DocxRenderer:
    """Turn one ``body.md`` into a Word document, honouring the report's route."""

    def __init__(self, config: ReportConfig) -> None:
        self.config = config
        self.document = Document()
        self.serif, self.sans = font_preferences(config)
        self.warnings: list[str] = []
        self.figure_number = 0
        self.heading_counters = [0, 0, 0]
        self.citation_numbers: dict[str, int] = {}
        self.bibliography: list[tuple[int, str, dict[str, str]]] = []
        self.bibliography_title = DEFAULT_BIBLIOGRAPHY_TITLE
        self._active_num_id: int | None = None

    # -- primitives --------------------------------------------------------

    @property
    def is_academic(self) -> bool:
        return self.config.route == "academic"

    def _add_segments(self, paragraph, segments: list[Segment], size_pt: float | None = None):
        for segment in segments:
            run = paragraph.add_run(segment.text)
            run.bold = segment.bold
            run.italic = segment.italic
            set_run_font(run, MONO_FONT if segment.mono else self.serif, size_pt)
        return paragraph

    def _paragraph(self, text: str, style: str | None = None, size_pt: float | None = None):
        paragraph = self.document.add_paragraph(style=style)
        self._add_segments(paragraph, inline_segments(text, self.citation_numbers), size_pt)
        return paragraph

    def add_page_break(self) -> None:
        self.document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # -- front matter ------------------------------------------------------

    def render_front_matter(self) -> None:
        """Route A gets the institutional cover; every other route a title block.

        ``references/document-routing.md`` is emphatic: routes B–E MUST NOT
        auto-include a UNL cover, teacher, subject, institutional motto or
        academic section numbering. A business deliverable therefore gets a
        sober title block and nothing else.
        """
        meta = self.config.metadata
        title = str(meta.get("title") or self.config.raw.get("title") or "Documento")
        if not self.is_academic:
            heading = self.document.add_paragraph(style="Title")
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_run_font(heading.add_run(title), self.serif, 20)
            byline = " — ".join(
                str(value)
                for value in (meta.get("student"), meta.get("date"))
                if str(value or "").strip()
            )
            if byline:
                line = self.document.add_paragraph(style=COVER_STYLE)
                line.alignment = WD_ALIGN_PARAGRAPH.LEFT
                set_run_font(line.add_run(byline), self.serif, 11)
            return

        if not self.config.academic_value("cover", "required", default=True):
            heading = self.document.add_paragraph(style="Title")
            set_run_font(heading.add_run(title), self.serif, 20)
            return

        self.render_academic_cover(meta, title)

    def render_academic_cover(self, meta: dict[str, Any], title: str) -> None:
        logo = ASSETS_DIR / LOGO_FILENAME
        if self.config.academic_value("cover", "logo_required", default=True) and logo.exists():
            picture_paragraph = self.document.add_paragraph(style=COVER_STYLE)
            picture_paragraph.add_run().add_picture(str(logo), width=Cm(4.0))

        university = str(meta.get("university") or DEFAULT_UNIVERSITY)
        faculty = str(meta.get("faculty") or DEFAULT_FACULTY)
        header = self.document.add_paragraph(style=COVER_STYLE)
        set_run_font(header.add_run(university), self.serif, 16)
        header.runs[0].bold = True
        faculty_paragraph = self.document.add_paragraph(style=COVER_STYLE)
        set_run_font(faculty_paragraph.add_run(faculty), self.serif, 12)

        self.document.add_paragraph(style=COVER_STYLE)
        title_paragraph = self.document.add_paragraph(style=COVER_STYLE)
        set_run_font(title_paragraph.add_run(title), self.serif, 18)
        title_paragraph.runs[0].bold = True
        activity = str(
            meta.get("activity") or self.config.raw.get("activity") or "Informe académico"
        )
        activity_paragraph = self.document.add_paragraph(style=COVER_STYLE)
        set_run_font(activity_paragraph.add_run(activity), self.serif, 12)
        self.document.add_paragraph(style=COVER_STYLE)

        rows = [
            (label, str(meta.get(key)).strip())
            for key, label in COVER_TABLE_FIELDS
            if str(meta.get(key) or "").strip()
        ]
        if rows:
            table = self.document.add_table(rows=0, cols=2)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for label, value in rows:
                cells = table.add_row().cells
                self._fill_cell(cells[0], label, bold=True, size_pt=11)
                self._fill_cell(cells[1], value, bold=False, size_pt=11)
        self.add_page_break()

    def _fill_cell(self, cell, text: str, *, bold: bool, size_pt: float) -> None:
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.first_line_indent = Cm(0)
        paragraph.paragraph_format.line_spacing = 1.0
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for segment in inline_segments(text, self.citation_numbers):
            run = paragraph.add_run(segment.text)
            run.bold = bold or segment.bold
            run.italic = segment.italic
            set_run_font(run, MONO_FONT if segment.mono else self.serif, size_pt)

    # -- block rendering ---------------------------------------------------

    def heading_prefix(self, level: int) -> str:
        """Academic section numbering — Route A machinery, nothing else."""
        if not self.is_academic or level > 3:
            return ""
        self.heading_counters[level - 1] += 1
        for deeper in range(level, 3):
            self.heading_counters[deeper] = 0
        return ".".join(str(n) for n in self.heading_counters[:level]) + "."

    def render_heading(self, level: int, title: str) -> None:
        paragraph = self.document.add_paragraph(style=f"Heading {level}")
        prefix = self.heading_prefix(level)
        if prefix:
            set_run_font(paragraph.add_run(f"{prefix} "), self.serif)
        self._add_segments(paragraph, inline_segments(title, self.citation_numbers))

    def render_code_block(self, lines: list[str]) -> None:
        block = list(lines)
        while block and not block[0].strip():
            block.pop(0)
        while block and not block[-1].strip():
            block.pop()
        if not block:
            return
        paragraph = self.document.add_paragraph(style=CODE_STYLE)
        set_paragraph_shading(paragraph, "F2F2F2")
        for index, line in enumerate(block):
            run = paragraph.add_run()
            if index:
                run.add_break()
            run.add_text(line.rstrip())
            set_run_font(run, MONO_FONT)

    def render_table(self, rows: list[list[str]]) -> None:
        if len(rows) < 2:
            return
        columns = max(len(row) for row in rows)
        table = self.document.add_table(rows=0, cols=columns)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for index, row in enumerate(rows):
            padded = row + [""] * (columns - len(row))
            cells = table.add_row().cells
            for cell, text in zip(cells, padded):
                self._fill_cell(cell, text, bold=index == 0, size_pt=10)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                if index == 0:
                    shading = OxmlElement("w:shd")
                    shading.set(qn("w:val"), "clear")
                    shading.set(qn("w:fill"), "EAEAEA")
                    cell._tc.get_or_add_tcPr().append(shading)
        repeat_table_header(table.rows[0])

    def render_figure(self, caption: str, source: str) -> None:
        if source.startswith(("http://", "https://", "data:")):
            self.warn(
                f"Figura remota omitida (Word no descarga imágenes al construir): {source}"
            )
            return
        resolved = resolve_figure(source, figure_search_dirs(self.config))
        if resolved is None:
            tried = ", ".join(str(path) for path in figure_search_dirs(self.config))
            raise SystemExit(
                f"Figura no encontrada: {source} (buscada en: {tried})"
            )
        picture_paragraph = self.document.add_paragraph()
        picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        picture_paragraph.paragraph_format.first_line_indent = Cm(0)
        picture = picture_paragraph.add_run().add_picture(str(resolved))
        max_width = self.text_width_emu()
        if picture.width > max_width:
            ratio = max_width / picture.width
            picture.width = int(picture.width * ratio)
            picture.height = int(picture.height * ratio)
        self.figure_number += 1
        caption_paragraph = self.document.add_paragraph(style=CAPTION_STYLE)
        set_run_font(caption_paragraph.add_run(f"Figura {self.figure_number}. "), self.sans, 10)
        for segment in inline_segments(caption, self.citation_numbers):
            run = caption_paragraph.add_run(segment.text)
            run.bold = segment.bold
            run.italic = segment.italic
            set_run_font(run, MONO_FONT if segment.mono else self.sans, 10)

    def text_width_emu(self) -> int:
        section = self.document.sections[0]
        return section.page_width - section.left_margin - section.right_margin

    def render_display_math(self, formula: str) -> None:
        paragraph = self.document.add_paragraph(style=MATH_STYLE)
        set_run_font(paragraph.add_run(formula.strip()), MONO_FONT)

    def render_bibliography(self) -> None:
        if not self.bibliography:
            return
        paragraph = self.document.add_paragraph(style="Heading 1")
        set_run_font(paragraph.add_run(self.bibliography_title), self.serif)
        for number, key, fields in self.bibliography:
            entry = self.document.add_paragraph(style=REFERENCE_STYLE)
            set_run_font(entry.add_run(format_reference(number, key, fields)), self.serif, 11)

    def warn(self, message: str) -> None:
        self.warnings.append(message)

    # -- driver ------------------------------------------------------------

    def render_body(self, markdown: str) -> None:
        lines = markdown.splitlines()
        paragraph_buffer: list[str] = []
        list_kind: str | None = None

        def flush_paragraph() -> None:
            nonlocal paragraph_buffer
            if paragraph_buffer:
                self._paragraph(" ".join(item.strip() for item in paragraph_buffer))
                paragraph_buffer = []

        def close_list() -> None:
            nonlocal list_kind
            list_kind = None
            self._active_num_id = None

        def open_list(kind: str) -> None:
            nonlocal list_kind
            if list_kind == kind:
                return
            close_list()
            list_kind = kind
            if kind == "ordered":
                self._active_num_id = new_numbering_id(self.document)

        def find_closing_fence(marker: str, start: int) -> int | None:
            for index in range(start, len(lines)):
                candidate = lines[index].strip()
                if len(candidate) >= 3 and set(candidate) == {marker}:
                    return index
            return None

        index = 0
        while index < len(lines):
            stripped = lines[index].rstrip().strip()
            if not stripped:
                flush_paragraph()
                close_list()
                index += 1
                continue

            fence = CODE_FENCE_RE.match(stripped)
            if fence:
                closing = find_closing_fence(fence.group("marker")[0], index + 1)
                if closing is None:
                    # Unterminated fence: drop the stray marker, keep parsing.
                    index += 1
                    continue
                flush_paragraph()
                close_list()
                self.render_code_block(lines[index + 1 : closing])
                index = closing + 1
                continue

            if "|" in stripped and index + 1 < len(lines) and is_table_separator(lines[index + 1]):
                flush_paragraph()
                close_list()
                rows = [split_table_row(stripped)]
                index += 2
                while index < len(lines) and lines[index].strip() and "|" in lines[index]:
                    rows.append(split_table_row(lines[index]))
                    index += 1
                self.render_table(rows)
                continue

            if stripped.startswith("finalanswer{") and stripped.endswith("}"):
                flush_paragraph()
                close_list()
                self.render_display_math(stripped[len("finalanswer{") : -1])
                index += 1
                continue

            display_math = DISPLAY_MATH_RE.match(stripped)
            if display_math:
                flush_paragraph()
                close_list()
                self.render_display_math(display_math.group("formula"))
                index += 1
                continue

            image = IMAGE_RE.match(stripped)
            if image:
                flush_paragraph()
                close_list()
                self.render_figure(image.group("caption"), image.group("src").strip().split(" ")[0])
                index += 1
                continue

            heading = HEADING_RE.match(stripped)
            if heading:
                flush_paragraph()
                close_list()
                title = heading.group(2).strip()
                if (
                    self.bibliography
                    and is_bibliography_heading(title)
                    and not any(rest.strip() for rest in lines[index + 1 :])
                ):
                    # This builder prints its own bibliography section below;
                    # keeping the author's heading would show the title twice.
                    self.bibliography_title = title
                    index += 1
                    continue
                self.render_heading(len(heading.group(1)), title)
                index += 1
                continue

            if stripped in PAGE_BREAK_COMMANDS or re.match(r"^-{3,}$", stripped):
                flush_paragraph()
                close_list()
                self.add_page_break()
                index += 1
                continue

            if LATEX_ENVIRONMENT_RE.match(stripped):
                flush_paragraph()
                close_list()
                self.warn(
                    "Comando LaTeX sin equivalente en Word, omitido para no imprimirlo "
                    f"como texto literal: {stripped}"
                )
                index += 1
                continue

            bullet = BULLET_RE.match(stripped)
            if bullet:
                flush_paragraph()
                open_list("bullet")
                self._paragraph(bullet.group(1), style="List Bullet")
                index += 1
                continue

            ordered = ORDERED_RE.match(stripped)
            if ordered:
                flush_paragraph()
                open_list("ordered")
                paragraph = self._paragraph(ordered.group(1), style="List Number")
                if self._active_num_id is not None:
                    bind_numbering(paragraph, self._active_num_id)
                index += 1
                continue

            paragraph_buffer.append(stripped)
            index += 1

        flush_paragraph()
        close_list()


def split_table_row(row: str) -> list[str]:
    return [cell.strip() for cell in row.strip().strip("|").split("|")]


def is_table_separator(row: str) -> bool:
    cells = split_table_row(row)
    return bool(cells) and all(re.match(r"^:?-{3,}:?$", cell) for cell in cells)


# ---------------------------------------------------------------------------
# Build
# ---------------------------------------------------------------------------


def render_document(config: ReportConfig) -> tuple[Document, list[str]]:
    """Render the configured report, returning the document and its warnings."""
    if not config.body_path.exists():
        raise SystemExit(f"No existe body.md: {config.body_path}")
    markdown = config.body_path.read_text(encoding="utf-8")
    bib_text = (
        config.bib_path.read_text(encoding="utf-8", errors="ignore")
        if config.bib_path
        else ""
    )

    renderer = DocxRenderer(config)
    numbers, bibliography, unused = resolve_citations(markdown, bib_text, config.bib_path)
    renderer.citation_numbers = numbers
    renderer.bibliography = bibliography
    if unused:
        renderer.warn("Entradas BibTeX no citadas: " + ", ".join(unused))

    math_spans = count_math_spans(markdown)
    if math_spans:
        renderer.warn(
            f"{math_spans} expresiones matemáticas insertadas como código fuente LaTeX: "
            "Word no tiene motor de ecuaciones y python-docx no genera OMML. "
            "Revisá las fórmulas antes de entregar, o usá el backend LaTeX para "
            "un documento con ecuaciones tipográficas."
        )

    apply_page_setup(renderer.document, config)
    apply_styles(renderer.document, config)
    declare_font_fallbacks(renderer.document, config)
    renderer.render_front_matter()
    renderer.render_body(markdown)
    renderer.render_bibliography()
    return renderer.document, renderer.warnings


def build(folder: Path) -> ReportConfig:
    config = load_report_config(folder)
    if config.backend != "docx":
        raise SystemExit(
            f"build_docx_report solo aplica a backend=docx; actual: {config.backend}"
        )
    document, warnings = render_document(config)
    config.docx_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(config.docx_path))
    for warning in warnings:
        print(f"Aviso: {warning}")
    print(f"DOCX generado: {config.docx_path}")
    if config.publish_global:
        published = publish_global_output(config.docx_path, config.metadata)
        if published:
            print(f"DOCX publicado por materia: {published}")
        else:
            print(
                "Aviso: no pude inferir la materia; no se publicó copia global en "
                "outputs/<materia>/"
            )
    return config


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("folder", type=Path, help="Carpeta del reporte con report.yml")
    args = parser.parse_args()
    build(args.folder)


if __name__ == "__main__":
    main()
