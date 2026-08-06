#!/usr/bin/env python3
"""Sequential quality-gated builder for the prompting IA Sistemas Operativos report.

Guarantees:
- stop on the first failed step
- keep intermediate DOCX/PDF/screenshots in backups, not outputs
- leave only the final PDF in outputs
- validate structure before delivering
"""
from __future__ import annotations

import argparse
import shutil
import subprocess
import sys
from datetime import datetime
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from report_config import CONTENT_ROOT

FONT = "Times New Roman"
ROOT = Path(__file__).resolve().parents[1]
# assets/ here holds the shipped logos plus the flow figure this script rewrites
# next to them, so it stays CODE. Deliverables and backups are content.
ASSETS = ROOT / "assets"
OUTPUTS = CONTENT_ROOT / "outputs"
BACKUPS = CONTENT_ROOT / "backups"
SO_OUTPUTS = OUTPUTS / "sistemas-operativos"
LOGO_RGB = ASSETS / "unl-logo-aa1.png"
LOGO_TRANSPARENT = ASSETS / "unl-logo-aa1-transparent.png"
FIGURE = ASSETS / "prompting-os-flow.png"
FINAL_NAME = "ensayo_prompting_ia_so_final.pdf"


def run(cmd: list[str], cwd: Path | None = None) -> subprocess.CompletedProcess[str]:
    print("$", " ".join(cmd))
    proc = subprocess.run(cmd, cwd=cwd, text=True, capture_output=True)
    if proc.returncode != 0:
        if proc.stdout:
            print(proc.stdout)
        if proc.stderr:
            print(proc.stderr, file=sys.stderr)
        raise SystemExit(f"FAILED: {' '.join(cmd)}")
    if proc.stdout.strip():
        print(proc.stdout.strip())
    return proc


def require(path: Path, label: str) -> None:
    if not path.exists():
        raise SystemExit(f"Missing {label}: {path}")


def set_run(run, size=None, bold=None, italic=None, color=None, font=FONT):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill="F2F2F2"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, val="single", sz="8", color="000000"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), val)
        node.set(qn("w:sz"), sz)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def clear_cell_borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "nil")


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, pct=10000):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(pct))
    tbl_w.set(qn("w:type"), "pct")


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def setup_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(12)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.first_line_indent = Cm(1.27)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.12
    for style_name, size in [("Heading 1", 14), ("Heading 2", 12)]:
        st = doc.styles[style_name]
        st.font.name = FONT
        st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor(0, 0, 0)
        st.paragraph_format.space_before = Pt(12)
        st.paragraph_format.space_after = Pt(6)
        st.paragraph_format.line_spacing = 1.0


def paragraph_border_bottom(p, size="18"):
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def create_transparent_logo() -> None:
    require(LOGO_RGB, "AA1 extracted RGB logo")
    # pdfimages produced RGB plus separate smask earlier. If smask is unavailable, make near-black transparent.
    smask_candidates = [Path("/tmp/aa1_extract/img-001.png"), ASSETS / "unl-logo-aa1-mask.png"]
    rgb = Image.open(LOGO_RGB).convert("RGBA")
    mask_path = next((p for p in smask_candidates if p.exists()), None)
    if mask_path:
        mask = Image.open(mask_path).convert("L").resize(rgb.size)
        rgb.putalpha(mask)
    else:
        # Fallback: remove black background from the extracted image.
        pix = rgb.load()
        for y in range(rgb.height):
            for x in range(rgb.width):
                r, g, b, a = pix[x, y]
                if r < 10 and g < 10 and b < 10:
                    pix[x, y] = (255, 255, 255, 0)
    rgb.save(LOGO_TRANSPARENT)
    require(LOGO_TRANSPARENT, "transparent logo")


def font_path() -> str:
    candidates = [
        "/usr/share/fonts/truetype/dejavu/DejaVuSerif.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSerif-Bold.ttf",
    ]
    for c in candidates:
        if Path(c).exists():
            return c
    return ""


def draw_wrapped(draw, text, xy, font, fill, max_width, line_gap=6):
    words = text.split()
    lines: list[str] = []
    line = ""
    for word in words:
        test = word if not line else f"{line} {word}"
        if draw.textbbox((0, 0), test, font=font)[2] <= max_width:
            line = test
        else:
            if line:
                lines.append(line)
            line = word
    if line:
        lines.append(line)
    x, y = xy
    for line in lines:
        draw.text((x, y), line, font=font, fill=fill)
        y += font.size + line_gap
    return y


def create_flow_figure() -> None:
    ASSETS.mkdir(parents=True, exist_ok=True)
    w, h = 1800, 760
    im = Image.new("RGB", (w, h), "white")
    d = ImageDraw.Draw(im)
    fp = font_path()
    title_font = ImageFont.truetype(fp, 44) if fp else ImageFont.load_default()
    box_font = ImageFont.truetype(fp, 30) if fp else ImageFont.load_default()
    small_font = ImageFont.truetype(fp, 25) if fp else ImageFont.load_default()
    d.text((w // 2, 45), "Flujo seguro de prompting con herramientas", font=title_font, fill="black", anchor="mm")

    boxes = [
        (90, 170, 350, 330, "Usuario\n+ prompt"),
        (440, 170, 700, 330, "Modelo\nde IA"),
        (790, 170, 1090, 330, "Capa de\nherramientas"),
        (1180, 170, 1460, 330, "Sistema\noperativo"),
        (1530, 140, 1740, 360, "Recursos:\narchivos\nred\nprocesos\nmemoria"),
    ]
    for x1, y1, x2, y2, text in boxes:
        d.rounded_rectangle((x1, y1, x2, y2), radius=18, outline="black", width=4, fill="#F2F2F2")
        lines = text.split("\n")
        total = len(lines) * 34
        y = y1 + ((y2 - y1) - total) // 2
        for line in lines:
            d.text(((x1 + x2) // 2, y), line, font=box_font, fill="black", anchor="ma")
            y += 38
    for box, next_box in zip(boxes, boxes[1:]):
        _, _, x2, _, _ = box
        nx1, _, _, _, _ = next_box
        y = 250
        d.line((x2 + 18, y, nx1 - 18, y), fill="black", width=4)
        d.polygon([(nx1 - 18, y), (nx1 - 38, y - 12), (nx1 - 38, y + 12)], fill="black")

    d.rounded_rectangle((260, 470, 1540, 650), radius=22, outline="black", width=3, fill="white")
    d.text((w // 2, 505), "Controles necesarios", font=box_font, fill="black", anchor="mm")
    controls = "permisos mínimos · sandbox/aislamiento · validación de salida · logs/auditoría · aprobación humana"
    draw_wrapped(d, controls, (345, 555), small_font, "black", 1110)
    im.save(FIGURE)
    require(FIGURE, "flow figure")


def add_center(doc, text="", size=12, bold=False, italic=False, after=3, before=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    set_run(r, size=size, bold=bold, italic=italic)
    return p


def add_header(section, logo_path: Path):
    header = section.header
    for p in header.paragraphs:
        p._element.getparent().remove(p._element)
    table = header.add_table(rows=1, cols=3, width=Inches(6.55))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Inches(1.42), Inches(4.10), Inches(1.03)]
    for idx, cell in enumerate(table.rows[0].cells):
        cell.width = widths[idx]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        clear_cell_borders(cell)
        set_cell_margins(cell, top=0, bottom=0, start=20, end=20)

    p = table.cell(0, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run().add_picture(str(logo_path), width=Inches(1.18))

    p = table.cell(0, 1).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, text in enumerate([
        "UNIVERSIDAD NACIONAL DE LOJA",
        "FACULTAD DE ENERGÍA, LAS INDUSTRIAS Y LOS RECURSOS",
        "NATURALES NO RENOVABLES",
        "CARRERA DE COMPUTACIÓN",
    ]):
        if i:
            p.add_run().add_break()
        set_run(p.add_run(text), size=7.15, bold=True)

    right = table.cell(0, 2)
    set_cell_borders(right, sz="8")
    p = right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, text in enumerate(["CARRERA DE", "COMPUTACIÓN", "HE-CIS-2022"]):
        if i:
            p.add_run().add_break()
        set_run(p.add_run(text), size=8.0, bold=True)

    line = header.add_paragraph()
    line.paragraph_format.space_before = Pt(1)
    line.paragraph_format.space_after = Pt(0)
    paragraph_border_bottom(line, size="18")


def add_footer(section):
    footer = section.footer
    for p in footer.paragraphs:
        p._element.getparent().remove(p._element)
    p = footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    set_run(p.add_run("Educamos para "), size=9.5, italic=True)
    set_run(p.add_run("Transformar"), size=9.5, italic=True, bold=True)


def add_cover(doc, meta, logo_path: Path):
    # AA1 cover: logo centered; internal pages: logo left in header.
    add_center(doc, "", after=8)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(logo_path), width=Inches(1.65))
    p.paragraph_format.space_after = Pt(18)

    add_center(doc, "UNIVERSIDAD NACIONAL DE LOJA", size=15, bold=True, after=4)
    add_center(doc, "FACULTAD DE LA ENERGÍA, LAS INDUSTRIAS Y", size=12.5, bold=True, after=0)
    add_center(doc, "LOS RECURSOS NATURALES NO RENOVABLES", size=12.5, bold=True, after=0)
    add_center(doc, "CARRERA COMPUTACIÓN", size=12.5, bold=True, after=22)
    add_center(doc, "Ensayo", size=20, bold=True, after=10)
    add_center(doc, meta.get("Título", "").upper(), size=12, bold=True, after=18)

    rows = [
        ("Asignatura", meta.get("Asignatura", "")),
        ("Título", meta.get("Título", "")),
        ("Tipo", meta.get("Tipo", "Ensayo")),
        ("Docente", meta.get("Docente", "")),
        ("Estudiante", meta.get("Estudiante", "")),
        ("Paralelo", meta.get("Paralelo", "")),
        ("Período\nAcadémico", meta.get("Período Académico", "")),
        ("Fecha", meta.get("Fecha", "")),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, (label, value) in enumerate(rows):
        c0, c1 = table.rows[i].cells
        c0.width = Inches(1.75)
        c1.width = Inches(4.15)
        for cell in (c0, c1):
            set_cell_borders(cell, sz="6")
            set_cell_margins(cell, top=85, bottom=85, start=120, end=120)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(c0, "F2F2F2")
        p0 = c0.paragraphs[0]
        for part_i, part in enumerate(label.split("\n")):
            if part_i:
                p0.add_run().add_break()
            set_run(p0.add_run(part), size=12, bold=True)
        set_run(c1.paragraphs[0].add_run(value), size=12)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(34)
    add_center(doc, 'Ciudad Universitaria "Guillermo Falconí Espinosa"', size=10.5, italic=True, after=0)
    add_center(doc, "Universidad Nacional de Loja", size=10.5, italic=True, after=0)
    doc.add_page_break()


def body_table_data():
    return [
        ("Técnica actual", "Aplicación y relación con Sistemas Operativos"),
        ("Prompt cero-disparo", "Pide una tarea sin ejemplos previos; reduce la interacción, pero puede producir salidas ambiguas que luego afecten scripts o archivos."),
        ("Few-shot prompting", "Usa ejemplos de entrada y salida; mejora la consistencia, aunque consume más contexto y memoria de sesión."),
        ("Prompt con rol y restricciones", "Define comportamiento, límites y formato; se parece a una política de ejecución para agentes con herramientas."),
        ("Descomposición de tareas", "Divide un problema grande en pasos verificables; permite revisar acciones antes de tocar procesos, comandos o archivos."),
        ("Prompt con herramientas", "Solicita uso de APIs, navegador, terminal o archivos; exige permisos mínimos, aislamiento, logs y auditoría."),
    ]


def add_body_table(doc):
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_width(table, 9300)
    for row_i, data in enumerate(body_table_data()):
        cells = table.add_row().cells
        for i, value in enumerate(data):
            cell = cells[i]
            cell.width = Inches(1.8 if i == 0 else 4.55)
            set_cell_borders(cell, sz="6")
            set_cell_margins(cell, top=65, bottom=65, start=75, end=75)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            if row_i == 0:
                set_cell_shading(cell, "EDEDED")
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            set_run(p.add_run(value), size=9.4, bold=True if row_i == 0 or i == 0 else False)
        set_row_cant_split(table.rows[-1])


def copy_paragraph(doc, src_p):
    text = src_p.text.strip()
    style = src_p.style.name
    if style == "Heading 1":
        p = doc.add_paragraph(style="Heading 1")
        set_run(p.add_run(text), size=14, bold=True, color="000000")
        p.paragraph_format.keep_with_next = True
        return p
    if style == "Heading 2":
        p = doc.add_paragraph(style="Heading 2")
        set_run(p.add_run(text), size=12, bold=True, color="000000")
        p.paragraph_format.keep_with_next = True
        return p

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.12
    if text.startswith("Tabla "):
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        set_run(p.add_run("Tabla 1. Técnicas actuales de prompting y relación con Sistemas Operativos."), size=10, italic=True)
        return p
    size = 10.5 if text.startswith("[") else 12
    if text.startswith("["):
        p.paragraph_format.left_indent = Cm(0.75)
        p.paragraph_format.first_line_indent = Cm(-0.75)
    for src_r in src_p.runs:
        set_run(p.add_run(src_r.text), size=size, bold=src_r.bold, italic=src_r.italic)
    if not src_p.runs and text:
        set_run(p.add_run(text), size=size)
    return p


def add_figure(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(FIGURE), width=Inches(5.85))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(6)
    set_run(cap.add_run("Figura 1. Flujo seguro de prompting con herramientas y recursos del sistema operativo."), size=10, italic=True)


def find_body_start(paragraphs):
    for i, p in enumerate(paragraphs):
        if p.style.name.startswith("Heading") and p.text.strip().startswith("1."):
            return i
    raise SystemExit("Could not find body start heading '1.'.")


def build_docx(source: Path, docx_out: Path) -> None:
    require(source, "source DOCX")
    create_transparent_logo()
    create_flow_figure()
    src = Document(source)
    if not src.tables:
        raise SystemExit("Expected metadata table in source DOCX.")
    meta = {row.cells[0].text.strip(): row.cells[1].text.strip() for row in src.tables[0].rows if len(row.cells) >= 2}

    doc = Document()
    setup_styles(doc)
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)
    sec.header_distance = Cm(0.8)
    sec.footer_distance = Cm(0.8)
    sec.different_first_page_header_footer = True
    add_header(sec, LOGO_TRANSPARENT)
    add_footer(sec)
    add_cover(doc, meta, LOGO_TRANSPARENT)

    body_start = find_body_start(src.paragraphs)
    inserted_table = False
    inserted_figure = False
    for src_p in src.paragraphs[body_start:]:
        text = src_p.text.strip()
        if not text:
            continue
        if text.startswith("Tabla ") and not inserted_table:
            doc.add_page_break()
            copy_paragraph(doc, src_p)
            add_body_table(doc)
            doc.add_paragraph()
            inserted_table = True
            continue
        copy_paragraph(doc, src_p)
        if text.startswith("Esta diferencia es clave") and not inserted_figure:
            add_figure(doc)
            inserted_figure = True

    docx_out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(docx_out)
    require(docx_out, "generated DOCX")


def validate_docx(docx_path: Path) -> None:
    doc = Document(docx_path)
    errors = []
    if doc.styles["Normal"].font.name != FONT:
        errors.append("Normal font is not Times New Roman")
    for name in ["Heading 1", "Heading 2"]:
        color = doc.styles[name].font.color.rgb
        if str(color) != "000000":
            errors.append(f"{name} is not black: {color}")
    if len(doc.inline_shapes) < 2:
        errors.append("Expected at least logo + figure images")
    if len(doc.tables) < 2:
        errors.append("Expected cover table + body table")
    if not doc.sections[0].different_first_page_header_footer:
        errors.append("Cover must not repeat body header/footer")
    if not doc.sections[0].header.tables:
        errors.append("Body header table missing")
    text = "\n".join(p.text for p in doc.paragraphs)
    for required in ["1. Tema", "2. Antecedentes", "Figura 1", "Tabla 1", "5. Bibliografía"]:
        if required not in text:
            errors.append(f"Missing required text: {required}")
    # Body table must be portrait-friendly: 2 columns, not 3+.
    if len(doc.tables) >= 2 and len(doc.tables[1].columns) != 2:
        errors.append("Body table must have 2 columns to avoid horizontal clipping")
    if errors:
        raise SystemExit("DOCX validation failed:\n- " + "\n- ".join(errors))


def convert_to_pdf(docx_path: Path, out_dir: Path) -> Path:
    run(["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", str(out_dir), str(docx_path)])
    pdf_path = out_dir / (docx_path.stem + ".pdf")
    require(pdf_path, "converted PDF")
    return pdf_path


def validate_pdf(pdf_path: Path, validation_dir: Path) -> None:
    info = run(["pdfinfo", str(pdf_path)]).stdout
    if "Pages:" not in info:
        raise SystemExit("PDF validation failed: page count unavailable")
    text = run(["pdftotext", "-layout", str(pdf_path), "-"]).stdout
    errors = []
    for required in ["UNIVERSIDAD NACIONAL DE LOJA", "1. Tema", "2. Antecedentes", "Tabla 1", "Figura 1", "5. Bibliografía"]:
        if required not in text:
            errors.append(f"PDF missing: {required}")
    if "Prompt cero-disparo" not in text or "permisos mínimos" not in text:
        errors.append("PDF missing compact body table content")
    if errors:
        raise SystemExit("PDF validation failed:\n- " + "\n- ".join(errors))
    validation_dir.mkdir(parents=True, exist_ok=True)
    run(["pdftoppm", "-png", "-f", "1", "-l", "3", "-r", "110", str(pdf_path), str(validation_dir / "page")])


def clean_outputs_keep_final(final_pdf: Path, backup_dir: Path) -> None:
    OUTPUTS.mkdir(parents=True, exist_ok=True)
    for item in list(OUTPUTS.rglob("*")):
        if item.is_dir() or item.name == ".gitkeep":
            continue
        if item.resolve() == final_pdf.resolve():
            continue
        dest = backup_dir / "previous_outputs" / item.relative_to(OUTPUTS)
        dest.parent.mkdir(parents=True, exist_ok=True)
        if dest.exists():
            dest.unlink()
        shutil.move(str(item), str(dest))


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--source", type=Path, default=OUTPUTS / "ensayo_prompting_ia_so.docx")
    parser.add_argument("--final-name", default=FINAL_NAME)
    args = parser.parse_args()

    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup_dir = BACKUPS / f"prompting_final_{ts}"
    staging = backup_dir / "staging"
    validation = backup_dir / "validation_pages"
    staging.mkdir(parents=True, exist_ok=True)

    print("STEP 1/7: validate inputs")
    if not args.source.exists():
        candidates = sorted(BACKUPS.glob("*/previous_outputs/ensayo_prompting_ia_so.docx"), key=lambda p: p.stat().st_mtime, reverse=True)
        if candidates:
            print(f"Source not in outputs; using latest backup source: {candidates[0]}")
            args.source = candidates[0]
    require(args.source, "source DOCX")
    require(LOGO_RGB, "AA1 logo RGB asset")

    print("STEP 2/7: build quality-gated DOCX")
    docx_out = staging / "ensayo_prompting_ia_so_final.docx"
    build_docx(args.source, docx_out)

    print("STEP 3/7: validate DOCX")
    validate_docx(docx_out)

    print("STEP 4/7: convert DOCX to PDF")
    pdf_staging = convert_to_pdf(docx_out, staging)

    print("STEP 5/7: validate PDF and generate review screenshots")
    validate_pdf(pdf_staging, validation)

    print("STEP 6/7: publish final PDF only")
    final_pdf = SO_OUTPUTS / args.final_name
    SO_OUTPUTS.mkdir(parents=True, exist_ok=True)
    if final_pdf.exists():
        shutil.move(str(final_pdf), str(backup_dir / final_pdf.name))
    shutil.copy2(pdf_staging, final_pdf)

    print("STEP 7/7: move previous outputs to backups")
    clean_outputs_keep_final(final_pdf, backup_dir)

    print(f"FINAL_PDF={final_pdf}")
    print(f"BACKUP_DIR={backup_dir}")


if __name__ == "__main__":
    main()
