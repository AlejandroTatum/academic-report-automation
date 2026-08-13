"""Behaviour tests for the generic body.md -> DOCX builder.

Every assertion runs against a REOPENED document (``Document(path)``) rather
than against the in-memory object the builder happened to return. A DOCX is a
zip of XML parts: the only proof that a paragraph style, a table cell or the
page geometry actually survived serialisation is reading it back off disk.
"""
from __future__ import annotations

import shutil
import sys
import zipfile
from pathlib import Path

import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm

TOOLS_DIR = str(Path(__file__).resolve().parent)
if TOOLS_DIR not in sys.path:
    sys.path.insert(0, TOOLS_DIR)

import build_docx_report as builder
from report_config import load_report_config
from validate_report import docx_validation

REPO_ROOT = Path(__file__).resolve().parents[1]
SAMPLE_IMAGE = REPO_ROOT / "assets" / "unl-logo-aa1-transparent.png"


# ---------------------------------------------------------------------------
# Fixture helpers
# ---------------------------------------------------------------------------


def make_report(
    tmp_path: Path,
    body: str,
    *,
    route: str | None = None,
    metadata: dict[str, str] | None = None,
    bib: str | None = None,
    extra: str = "",
) -> Path:
    """Write a minimal report folder and return it."""
    folder = tmp_path / "report_docx"
    folder.mkdir(parents=True, exist_ok=True)
    (folder / "body.md").write_text(body, encoding="utf-8")
    meta = metadata or {
        "title": "Documento de prueba",
        "subject": "Sistemas Operativos",
        "teacher": "Ing. Prueba",
        "student": "Alejandro Padilla",
        "date": "7 de agosto de 2026",
    }
    lines = ["type: docx", "backend: docx", "output: docx", "publish_global: false"]
    if route:
        lines.append(f"route: {route}")
    lines.append("metadata:")
    lines.extend(f'  {key}: "{value}"' for key, value in meta.items())
    if bib is not None:
        (folder / "sources.bib").write_text(bib, encoding="utf-8")
    if extra:
        lines.append(extra)
    (folder / "report.yml").write_text("\n".join(lines) + "\n", encoding="utf-8")
    return folder


def build_and_open(folder: Path) -> Document:
    config = builder.build(folder)
    assert config.docx_path.exists(), f"el builder no escribió {config.docx_path}"
    return Document(str(config.docx_path))


def paragraph_texts(document: Document) -> list[str]:
    return [p.text for p in document.paragraphs]


def styles_used(document: Document) -> list[str]:
    return [p.style.name for p in document.paragraphs]


def start_override_values(document: Document) -> dict[int, int]:
    """Map each ``w:numId`` to its explicit ``w:startOverride`` value, if any.

    A cloned numbering definition restarts its list at the override value; a
    plain clone that merely reuses the shared abstract definition has none.
    """
    numbering = document.part.numbering_part.element
    overrides: dict[int, int] = {}
    for num in numbering.findall(qn("w:num")):
        num_id = num.get(qn("w:numId"))
        if num_id is None:
            continue
        for level in num.findall(qn("w:lvlOverride")):
            for start in level.findall(qn("w:startOverride")):
                value = start.get(qn("w:val"))
                if value is not None:
                    overrides[int(num_id)] = int(value)
    return overrides


# ---------------------------------------------------------------------------
# Page setup and typography — the academic_format.yml contract
# ---------------------------------------------------------------------------


def test_page_is_a4_with_configured_margins(tmp_path: Path) -> None:
    document = build_and_open(make_report(tmp_path, "# Título\n\nTexto.\n"))
    section = document.sections[0]
    assert abs(section.page_width - Cm(21.0)) < Cm(0.05)
    assert abs(section.page_height - Cm(29.7)) < Cm(0.05)
    for margin in (
        section.top_margin,
        section.bottom_margin,
        section.left_margin,
        section.right_margin,
    ):
        assert abs(margin - Cm(2.5)) < Cm(0.05)


def test_body_typography_matches_format_contract(tmp_path: Path) -> None:
    document = build_and_open(make_report(tmp_path, "# Título\n\nTexto del cuerpo.\n"))
    normal = document.styles["Normal"]
    assert normal.font.name == "Times New Roman"
    assert normal.font.size.pt == 12
    assert normal.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY
    assert abs(normal.paragraph_format.line_spacing - 1.15) < 0.001
    assert abs(normal.paragraph_format.first_line_indent - Cm(1.27)) < Cm(0.02)


def test_font_fallback_chain_is_declared_in_the_font_table(tmp_path: Path) -> None:
    """Word/LibreOffice substitute through ``w:altName``, not through CSS-style lists."""
    folder = make_report(tmp_path, "# Título\n\nTexto.\n")
    config = builder.build(folder)
    with zipfile.ZipFile(config.docx_path) as archive:
        font_table = archive.read("word/fontTable.xml").decode("utf-8")
    assert "TeX Gyre Termes" in font_table
    assert "TeX Gyre Heros" in font_table


# ---------------------------------------------------------------------------
# Block grammar — the same documents markdown_to_latex() accepts
# ---------------------------------------------------------------------------


def test_headings_map_to_heading_levels_one_to_four(tmp_path: Path) -> None:
    body = "# Uno\n\n## Dos\n\n### Tres\n\n#### Cuatro\n\nTexto.\n"
    document = build_and_open(make_report(tmp_path, body, route="technical"))
    used = styles_used(document)
    for level in (1, 2, 3, 4):
        assert f"Heading {level}" in used


def test_inline_bold_italic_and_code(tmp_path: Path) -> None:
    body = "# T\n\nUn **negrita**, una *cursiva* y `codigo_inline`.\n"
    document = build_and_open(make_report(tmp_path, body, route="technical"))
    runs = [run for p in document.paragraphs for run in p.runs]
    assert any(run.bold and run.text == "negrita" for run in runs)
    assert any(run.italic and run.text == "cursiva" for run in runs)
    assert any(run.text == "codigo_inline" and run.font.name == builder.MONO_FONT for run in runs)


# ---------------------------------------------------------------------------
# Nested inline semantics (#6) — shared DOCX/LaTeX fixture
# ---------------------------------------------------------------------------

# Each entry: (markdown, inner text, expected flags on the inner run, LaTeX fragment).
# Both backends must render the same visible text with no leftover markers:
# DOCX segments carry the outer flag merged with the inner rule's mono flag,
# and LaTeX nests \texttt / math / \cite inside the emphasis command.
NESTED_INLINE_FIXTURES = [
    (
        "**Label (`fieldName`, extra):** value.",
        "fieldName",
        {"bold": True, "mono": True},
        r"\texttt{fieldName}",
    ),
    (
        "**Energía $E = mc^2$ total**",
        "E = mc^2",
        {"bold": True, "mono": True},
        r"$E = mc^2$",
    ),
    (
        "*El archivo `config.yml` manda*",
        "config.yml",
        {"italic": True, "mono": True},
        r"\texttt{config.yml}",
    ),
    (
        "**Ver [@torres2024]**",
        "[1]",
        {"bold": True},
        r"\cite{torres2024}",
    ),
]


def test_inline_code_inside_bold_merges_mono_flag() -> None:
    """Inline code inside a bold span must keep both flags and drop the backticks."""
    segments = builder.inline_segments("**Label (`fieldName`, extra):** value.")
    assert "".join(segment.text for segment in segments) == "Label (fieldName, extra): value."
    assert not any("`" in segment.text for segment in segments)
    code_segments = [segment for segment in segments if segment.text == "fieldName"]
    assert code_segments, "el código anidado debe ser un segmento propio, no backticks literales"
    assert code_segments[0].bold is True
    assert code_segments[0].mono is True


@pytest.mark.parametrize(
    "markdown,inner_text,expected_flags,_latex_fragment",
    NESTED_INLINE_FIXTURES,
)
def test_nested_inline_segments_merge_outer_and_inner_flags(
    markdown: str, inner_text: str, expected_flags: dict[str, bool], _latex_fragment: str
) -> None:
    segments = builder.inline_segments(markdown, {"torres2024": 1})
    assert "".join(segment.text for segment in segments) != markdown
    assert not any(
        marker in segment.text for segment in segments for marker in ("`", "$", "[@", "*")
    )
    inner = [segment for segment in segments if segment.text == inner_text]
    assert inner, f"segmento interno {inner_text!r} ausente en {markdown!r}"
    segment = inner[0]
    assert segment.bold is expected_flags.get("bold", False)
    assert segment.italic is expected_flags.get("italic", False)
    assert segment.mono is expected_flags.get("mono", False)


def test_nested_inline_docx_and_latex_agree_on_shared_fixture() -> None:
    """Both backends render the same fixture with no leftover Markdown markers."""
    import build_latex_report

    for markdown, _inner_text, _expected_flags, latex_fragment in NESTED_INLINE_FIXTURES:
        segments = builder.inline_segments(markdown, {"torres2024": 1})
        assert not any(
            marker in segment.text for segment in segments for marker in ("`", "$", "[@", "*")
        )
        latex = build_latex_report.convert_inline(markdown)
        assert "`" not in latex
        assert "[@" not in latex
        assert "**" not in latex
        # \allowbreak{} is a zero-width line-break penalty, not a glyph; strip
        # it so the fragment check reads the visually rendered LaTeX.
        assert latex_fragment in latex.replace(r"\allowbreak{}", "")


def test_bullet_and_ordered_lists(tmp_path: Path) -> None:
    body = "# T\n\n- alfa\n- beta\n\nTexto.\n\n1. primero\n2. segundo\n"
    document = build_and_open(make_report(tmp_path, body, route="technical"))
    pairs = [(p.style.name, p.text) for p in document.paragraphs]
    assert ("List Bullet", "alfa") in pairs
    assert ("List Bullet", "beta") in pairs
    assert ("List Number", "primero") in pairs
    assert ("List Number", "segundo") in pairs


def test_ordered_lists_restart_their_numbering(tmp_path: Path) -> None:
    """Two procedures must both start at 1, not continue 1..4.

    Distinct ``w:numId`` values alone do not restart a list: both clones point
    at the same abstract definition, so Word keeps counting unless each clone
    carries an explicit ``w:lvlOverride``/``w:startOverride val="1"``.
    """
    body = "# T\n\n1. uno\n2. dos\n\nTexto intermedio.\n\n1) otro uno\n2) otro dos\n"
    document = build_and_open(make_report(tmp_path, body, route="technical"))
    num_ids = [
        builder.paragraph_num_id(p)
        for p in document.paragraphs
        if p.style.name == "List Number"
    ]
    assert len(num_ids) == 4
    assert all(value is not None for value in num_ids)
    assert num_ids[0] == num_ids[1]
    assert num_ids[2] == num_ids[3]
    assert num_ids[0] != num_ids[2]
    # Each cloned list definition must restart its own counter at 1.
    overrides = start_override_values(document)
    assert set(overrides) == set(num_ids)
    assert all(value == 1 for value in overrides.values())


def test_pipe_table_keeps_header_and_cells(tmp_path: Path) -> None:
    body = (
        "# T\n\n"
        "| Elemento | Fuente | Tamaño |\n"
        "| --- | --- | --- |\n"
        "| Cuerpo | Times New Roman | 12 pt |\n"
        "| Pie | Arial | 10 pt |\n"
    )
    document = build_and_open(make_report(tmp_path, body, route="technical"))
    assert len(document.tables) == 1
    table = document.tables[0]
    assert len(table.rows) == 3
    assert len(table.columns) == 3
    assert [cell.text for cell in table.rows[0].cells] == ["Elemento", "Fuente", "Tamaño"]
    assert [cell.text for cell in table.rows[2].cells] == ["Pie", "Arial", "10 pt"]
    header_runs = [run for cell in table.rows[0].cells for p in cell.paragraphs for run in p.runs]
    assert header_runs and all(run.bold for run in header_runs)


def test_ragged_table_row_is_padded_not_dropped(tmp_path: Path) -> None:
    body = "# T\n\n| A | B | C |\n| --- | --- | --- |\n| solo |\n"
    document = build_and_open(make_report(tmp_path, body, route="technical"))
    table = document.tables[0]
    assert [cell.text for cell in table.rows[1].cells] == ["solo", "", ""]


@pytest.mark.parametrize("fence", ["```", "~~~"])
def test_code_fence_preserves_lines_and_indentation(tmp_path: Path, fence: str) -> None:
    body = f"# T\n\n{fence}python\ndef f():\n    return 1\n{fence}\n"
    document = build_and_open(make_report(tmp_path, body, route="technical"))
    blocks = [p for p in document.paragraphs if p.style.name == builder.CODE_STYLE]
    assert len(blocks) == 1
    text = blocks[0].text
    assert "def f():" in text
    assert "    return 1" in text
    # The language word labels the fence, it is not part of the source.
    assert "python" not in text
    assert all(run.font.name == builder.MONO_FONT for run in blocks[0].runs if run.text)


def test_unterminated_code_fence_does_not_swallow_the_document(tmp_path: Path) -> None:
    body = "# T\n\n```\nsin cierre\n\n# Sigue\n\nTexto final.\n"
    document = build_and_open(make_report(tmp_path, body, route="technical"))
    assert "Sigue" in " ".join(paragraph_texts(document))


def test_horizontal_rule_and_newpage_become_page_breaks(tmp_path: Path) -> None:
    body = "# Uno\n\nA.\n\n---\n\n# Dos\n\nB.\n\n\\newpage\n\n# Tres\n\nC.\n"
    folder = make_report(tmp_path, body, route="technical")
    config = builder.build(folder)
    with zipfile.ZipFile(config.docx_path) as archive:
        xml = archive.read("word/document.xml").decode("utf-8")
    assert xml.count('w:type="page"') == 2


# ---------------------------------------------------------------------------
# Figures
# ---------------------------------------------------------------------------


def test_image_is_embedded_with_a_caption(tmp_path: Path) -> None:
    folder = make_report(
        tmp_path,
        "# T\n\n![Diagrama de flujo](../../figura.png)\n",
        route="technical",
    )
    # The LaTeX branch resolves figure references from <report>/build/, which is
    # why real bodies carry ../../../assets/generated/... paths.
    shutil.copy2(SAMPLE_IMAGE, tmp_path / "figura.png")
    document = Document(str(builder.build(folder).docx_path))
    assert len(document.inline_shapes) == 1
    captions = [p.text for p in document.paragraphs if p.style.name == builder.CAPTION_STYLE]
    assert captions == ["Figura 1. Diagrama de flujo"]


def test_missing_image_fails_naming_the_path(tmp_path: Path) -> None:
    folder = make_report(tmp_path, "# T\n\n![Falta](assets/no_existe.png)\n", route="technical")
    with pytest.raises(SystemExit) as excinfo:
        builder.build(folder)
    message = str(excinfo.value)
    assert "assets/no_existe.png" in message
    assert "Figura no encontrada" in message


def test_remote_image_is_skipped_with_a_warning(tmp_path: Path, capsys) -> None:
    folder = make_report(
        tmp_path, "# T\n\n![Remota](https://ejemplo.org/x.png)\n", route="technical"
    )
    builder.build(folder)
    assert "https://ejemplo.org/x.png" in capsys.readouterr().out


# ---------------------------------------------------------------------------
# Route separation — document-routing.md is binding
# ---------------------------------------------------------------------------


def test_academic_route_renders_cover_and_metadata_table(tmp_path: Path) -> None:
    folder = make_report(
        tmp_path,
        "# Desarrollo\n\nTexto.\n",
        metadata={
            "title": "Planificación de CPU",
            "subject": "Sistemas Operativos",
            "teacher": "Ing. Hernán Torres",
            "student": "Alejandro Padilla",
            "parallel": "A",
            "career": "Carrera de Computación",
            "date": "7 de agosto de 2026",
        },
    )
    document = Document(str(builder.build(folder).docx_path))
    text = "\n".join(paragraph_texts(document))
    assert "Universidad Nacional de Loja" in text
    assert document.tables, "la ruta académica debe traer la tabla de metadatos"
    cover_table = "\n".join(cell.text for row in document.tables[0].rows for cell in row.cells)
    assert "Ing. Hernán Torres" in cover_table
    assert "Sistemas Operativos" in cover_table
    assert "Carrera de Computación" in cover_table
    # Academic section numbering is Route A machinery.
    assert any(p.text == "1. Desarrollo" for p in document.paragraphs)


@pytest.mark.parametrize("route", ["project", "business", "technical", "other"])
def test_non_academic_routes_get_no_institutional_furniture(tmp_path: Path, route: str) -> None:
    folder = make_report(
        tmp_path,
        "# Propósito\n\nTexto.\n",
        route=route,
        metadata={
            "title": "Contrato del router",
            "student": "Plataforma de documentos",
            "date": "7 de agosto de 2026",
        },
    )
    document = Document(str(builder.build(folder).docx_path))
    text = "\n".join(paragraph_texts(document))
    assert "Universidad Nacional de Loja" not in text
    assert "Docente" not in text
    assert "Asignatura" not in text
    assert document.tables == []
    # Title block only: the title and the author must still be there.
    assert "Contrato del router" in text
    assert "Plataforma de documentos" in text
    # No academic section numbering.
    assert any(p.text == "Propósito" for p in document.paragraphs)
    assert not any(p.text.startswith("1. ") for p in document.paragraphs)


def test_academic_cover_can_be_relaxed_per_report(tmp_path: Path) -> None:
    folder = make_report(
        tmp_path,
        "# Desarrollo\n\nTexto.\n",
        extra="cover:\n  required: false\n  logo_required: false",
    )
    document = Document(str(builder.build(folder).docx_path))
    text = "\n".join(paragraph_texts(document))
    assert "Universidad Nacional de Loja" not in text
    assert document.tables == []


# ---------------------------------------------------------------------------
# Citations and bibliography
# ---------------------------------------------------------------------------


BIB = """
@misc{padilla2026,
  author = {Padilla Espinoza, Alejandro Emanuel},
  title = {Academic Report Automation Toolkit},
  year = {2026},
  howpublished = {\\url{https://github.com/example}}
}
@article{torres2024,
  author = {Torres, Hernán},
  title = {Planificación de CPU},
  journal = {Revista de Sistemas},
  year = {2024},
  doi = {10.1000/xyz}
}
"""


def test_citations_are_numbered_in_order_of_first_appearance(tmp_path: Path) -> None:
    body = (
        "# T\n\nPrimero [@torres2024] y luego [@padilla2026].\n\n"
        "Vuelve a citar [@torres2024].\n\n# Referencias\n"
    )
    folder = make_report(tmp_path, body, route="technical", bib=BIB)
    document = Document(str(builder.build(folder).docx_path))
    text = "\n".join(paragraph_texts(document))
    assert "Primero [1] y luego [2]." in text
    assert "Vuelve a citar [1]." in text
    assert "[@" not in text
    # A bibliography section, printed once, listing both cited works.
    assert text.count("Referencias") == 1
    entries = [p.text for p in document.paragraphs if p.style.name == builder.REFERENCE_STYLE]
    assert len(entries) == 2
    assert entries[0].startswith("[1]")
    assert "Torres" in entries[0]
    assert entries[1].startswith("[2]")
    assert "Padilla" in entries[1]


def test_multi_key_citation_expands_to_every_number(tmp_path: Path) -> None:
    body = "# T\n\nAmbos [@torres2024, padilla2026].\n"
    folder = make_report(tmp_path, body, route="technical", bib=BIB)
    document = Document(str(builder.build(folder).docx_path))
    assert "Ambos [1], [2]." in "\n".join(paragraph_texts(document))


def test_citation_without_bib_entry_is_a_hard_error(tmp_path: Path) -> None:
    body = "# T\n\nCita rota [@fantasma].\n"
    folder = make_report(tmp_path, body, route="technical", bib=BIB)
    with pytest.raises(SystemExit) as excinfo:
        builder.build(folder)
    assert "fantasma" in str(excinfo.value)


def test_citation_without_any_bibliography_is_a_hard_error(tmp_path: Path) -> None:
    folder = make_report(tmp_path, "# T\n\nCita [@suelta].\n", route="technical")
    with pytest.raises(SystemExit) as excinfo:
        builder.build(folder)
    assert "suelta" in str(excinfo.value)


def test_uncited_bib_entry_is_reported_not_silently_dropped(tmp_path: Path, capsys) -> None:
    body = "# T\n\nSolo una [@torres2024].\n"
    folder = make_report(tmp_path, body, route="technical", bib=BIB)
    builder.build(folder)
    assert "padilla2026" in capsys.readouterr().out


def test_body_without_citations_gets_no_bibliography(tmp_path: Path) -> None:
    folder = make_report(tmp_path, "# T\n\nSin citas.\n", route="technical", bib=BIB)
    document = Document(str(builder.build(folder).docx_path))
    assert not [p for p in document.paragraphs if p.style.name == builder.REFERENCE_STYLE]


# ---------------------------------------------------------------------------
# Math
# ---------------------------------------------------------------------------


def test_math_is_rendered_as_source_in_a_distinct_style_and_warned(tmp_path: Path, capsys) -> None:
    body = "# T\n\nSea $E = mc^2$ la energía.\n\n$$\\int_0^1 x\\,dx = \\frac{1}{2}$$\n"
    folder = make_report(tmp_path, body, route="technical")
    document = Document(str(builder.build(folder).docx_path))
    out = capsys.readouterr().out
    assert "ecuaciones" in out.lower() or "matem" in out.lower()

    display = [p for p in document.paragraphs if p.style.name == builder.MATH_STYLE]
    assert len(display) == 1
    assert "\\int_0^1" in display[0].text

    inline_runs = [
        run
        for p in document.paragraphs
        for run in p.runs
        if run.text.strip() == "E = mc^2"
    ]
    assert inline_runs, "la matemática inline debe conservarse legible"
    assert inline_runs[0].font.name == builder.MONO_FONT
    # The source is preserved verbatim, never half-escaped into prose.
    assert "$" not in "\n".join(paragraph_texts(document))


# ---------------------------------------------------------------------------
# Wiring
# ---------------------------------------------------------------------------


def test_build_refuses_a_non_docx_backend(tmp_path: Path) -> None:
    folder = tmp_path / "latex_report"
    folder.mkdir()
    (folder / "body.md").write_text("# T\n", encoding="utf-8")
    (folder / "report.yml").write_text(
        'backend: latex\nmetadata:\n  title: "T"\n  subject: "S"\n'
        '  teacher: "D"\n  student: "A"\n  date: "hoy"\n',
        encoding="utf-8",
    )
    with pytest.raises(SystemExit) as excinfo:
        builder.build(folder)
    assert "docx" in str(excinfo.value)


def test_missing_body_fails_with_a_spanish_message(tmp_path: Path) -> None:
    folder = make_report(tmp_path, "# T\n")
    (folder / "body.md").unlink()
    with pytest.raises(SystemExit) as excinfo:
        builder.build(folder)
    assert "No existe body.md" in str(excinfo.value)


def test_publish_global_is_called_when_configured(tmp_path: Path, monkeypatch) -> None:
    published: list[Path] = []
    monkeypatch.setattr(
        builder,
        "publish_global_output",
        lambda source, metadata=None: published.append(Path(source)) or Path(source),
    )
    folder = make_report(tmp_path, "# T\n\nTexto.\n")
    (folder / "report.yml").write_text(
        (folder / "report.yml").read_text(encoding="utf-8").replace(
            "publish_global: false", "publish_global: true"
        ),
        encoding="utf-8",
    )
    builder.build(folder)
    assert published


def test_router_dispatches_the_docx_backend() -> None:
    import build_report_auto

    assert "docx" not in build_report_auto.UNSUPPORTED_BUILDER_MESSAGES
    assert "visual" in build_report_auto.UNSUPPORTED_BUILDER_MESSAGES


# ---------------------------------------------------------------------------
# docx_validation() against a real artefact
# ---------------------------------------------------------------------------


def test_docx_validation_passes_for_a_generated_document(tmp_path: Path) -> None:
    body = (
        "# Uno\n\nTexto suficiente para no parecer vacío.\n\n"
        "| A | B |\n| --- | --- |\n| 1 | 2 |\n\n"
        "![Figura](../../figura.png)\n"
    )
    folder = make_report(tmp_path, body)
    shutil.copy2(SAMPLE_IMAGE, tmp_path / "figura.png")
    builder.build(folder)
    result = docx_validation(load_report_config(folder))
    assert result.errors == []


def test_docx_validation_reports_a_missing_file(tmp_path: Path) -> None:
    folder = make_report(tmp_path, "# T\n\nTexto.\n")
    result = docx_validation(load_report_config(folder))
    assert any("No existe DOCX" in error for error in result.errors)


def test_docx_validation_rejects_an_empty_document(tmp_path: Path) -> None:
    folder = make_report(tmp_path, "# T\n\nTexto.\n")
    config = load_report_config(folder)
    config.docx_path.parent.mkdir(parents=True, exist_ok=True)
    Document().save(str(config.docx_path))
    result = docx_validation(config)
    assert any("sin contenido" in error.lower() for error in result.errors)


def test_docx_validation_rejects_a_corrupt_file(tmp_path: Path) -> None:
    folder = make_report(tmp_path, "# T\n\nTexto.\n")
    config = load_report_config(folder)
    config.docx_path.parent.mkdir(parents=True, exist_ok=True)
    config.docx_path.write_bytes(b"esto no es un docx" * 800)
    result = docx_validation(config)
    assert any("no se pudo abrir" in error.lower() for error in result.errors)


def test_docx_validation_rejects_wrong_page_size(tmp_path: Path) -> None:
    folder = make_report(tmp_path, "# T\n\nTexto suficiente para el documento.\n")
    config = load_report_config(folder)
    builder.build(folder)
    document = Document(str(config.docx_path))
    document.sections[0].page_width = Cm(21.59)
    document.sections[0].page_height = Cm(27.94)
    document.save(str(config.docx_path))
    result = docx_validation(config)
    assert any("A4" in error for error in result.errors)


def test_docx_validation_rejects_a_document_missing_its_figures(tmp_path: Path) -> None:
    body = "# T\n\nTexto.\n\n![Figura](../../figura.png)\n"
    folder = make_report(tmp_path, body)
    shutil.copy2(SAMPLE_IMAGE, tmp_path / "figura.png")
    config = load_report_config(folder)
    builder.build(folder)
    # Rebuild a picture-free document at the same path.
    stripped = Document()
    stripped.add_paragraph("Texto largo suficiente para superar el umbral vacío.")
    stripped.add_heading("Uno", level=1)
    stripped.save(str(config.docx_path))
    result = docx_validation(config)
    assert any("imágenes" in error for error in result.errors)
