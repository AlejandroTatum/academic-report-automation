"""Delivery-folder verification tests (issue #9 clean-delivery).

``verify_delivery()`` proves a run's delivery folder holds only clean finals:
every expected PDF/DOCX exists, is non-empty, is readable, and carries a
recorded SHA-256 and page count, and no stray working artifact leaked in.
PDF fixtures are real multi-page files written by PIL (the visual auditor's
stack); page counts come from ``pdfinfo``.
"""
from __future__ import annotations

import hashlib
import shutil
import sys
from pathlib import Path

import pytest
from docx import Document
from docx.enum.text import WD_BREAK
from PIL import Image

TOOLS_DIR = str(Path(__file__).resolve().parent)
if TOOLS_DIR not in sys.path:
    sys.path.insert(0, TOOLS_DIR)

from verify_delivery import verify_delivery  # noqa: E402

pytestmark = pytest.mark.skipif(
    shutil.which("pdfinfo") is None, reason="pdfinfo (poppler-utils) not installed",
)


def delivery_folder(tmp_path: Path) -> Path:
    folder = tmp_path / "entregas"
    folder.mkdir()
    return folder


def _write_pdf(path: Path, pages: int) -> None:
    images = [Image.new("RGB", (2480, 3508), "white") for _ in range(pages)]
    images[0].save(path, save_all=True, append_images=images[1:], resolution=150.0)


def _write_docx(path: Path, page_breaks: int) -> None:
    document = Document()
    for _ in range(page_breaks):
        document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    document.save(path)


# -- Spec scenario: clean run passes with full evidence ----------------------


def test_delivery_final_artifact_verified(tmp_path) -> None:
    """A clean run records existence, non-emptiness, readability, hash, pages."""
    folder = delivery_folder(tmp_path)
    final = folder / "informe.pdf"
    _write_pdf(final, pages=3)
    raw = final.read_bytes()

    result = verify_delivery(folder, [final])

    assert result.ok is True
    assert result.errors == []
    evidence = result.files[0]
    assert evidence.exists is True
    assert evidence.readable is True
    assert evidence.size_bytes == len(raw) and evidence.size_bytes > 0
    assert evidence.sha256 == hashlib.sha256(raw).hexdigest()
    assert evidence.page_count == 3
    assert result.stray == []


# -- Spec scenario: non-final artifact is rejected ---------------------------


def test_non_final_artifact_rejected_from_delivery(tmp_path) -> None:
    """Working evidence in the delivery folder fails the run, naming the file."""
    folder = delivery_folder(tmp_path)
    final = folder / "informe.pdf"
    _write_pdf(final, pages=1)
    stray = folder / "report.yml"
    stray.write_text("title: Informe\n", encoding="utf-8")

    result = verify_delivery(folder, [final])

    assert result.ok is False
    assert result.stray == [stray]
    assert any("report.yml" in error for error in result.errors)


# -- Triangulation: second artifact type, second code path -------------------


def test_docx_final_reports_page_count_from_page_breaks(tmp_path) -> None:
    """A DOCX final records pagination from its explicit page breaks."""
    folder = delivery_folder(tmp_path)
    final = folder / "informe.docx"
    _write_docx(final, page_breaks=1)

    result = verify_delivery(folder, [final])

    assert result.ok is True
    evidence = result.files[0]
    assert evidence.page_count == 2
    assert evidence.sha256 == hashlib.sha256(final.read_bytes()).hexdigest()


def test_two_finals_are_both_verified(tmp_path) -> None:
    """A run with two expected finals verifies each independently."""
    folder = delivery_folder(tmp_path)
    first = folder / "capitulo-a.pdf"
    second = folder / "capitulo-b.pdf"
    _write_pdf(first, pages=1)
    _write_pdf(second, pages=2)

    result = verify_delivery(folder, [first, second])

    assert result.ok is True
    assert [evidence.page_count for evidence in result.files] == [1, 2]


# -- Triangulation: gate failures -------------------------------------------


def test_missing_expected_final_fails(tmp_path) -> None:
    """An expected final that was never copied into the folder fails the run."""
    folder = delivery_folder(tmp_path)
    missing = folder / "informe.pdf"

    result = verify_delivery(folder, [missing])

    assert result.ok is False
    assert any("informe.pdf" in error for error in result.errors)
    assert any("missing" in error.lower() for error in result.errors)


def test_empty_final_fails(tmp_path) -> None:
    """A zero-byte final fails the non-emptiness gate."""
    folder = delivery_folder(tmp_path)
    empty = folder / "informe.pdf"
    empty.write_bytes(b"")

    result = verify_delivery(folder, [empty])

    assert result.ok is False
    assert any("empty" in error.lower() for error in result.errors)


def test_directory_expected_final_is_not_readable(tmp_path) -> None:
    """An expected path that cannot be read as a file fails the readability gate."""
    folder = delivery_folder(tmp_path)
    not_a_file = folder / "informe.pdf"
    not_a_file.mkdir()

    result = verify_delivery(folder, [not_a_file])

    assert result.ok is False
    assert any("read" in error.lower() for error in result.errors)
