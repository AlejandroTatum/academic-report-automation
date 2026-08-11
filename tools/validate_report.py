#!/usr/bin/env python3
"""Common, backend-aware validation layer for Alejandro's university reports."""
from __future__ import annotations

import argparse
import os
import re
import subprocess
from dataclasses import dataclass, field
from pathlib import Path

from report_config import ROOT as AUTOMATION_ROOT
from report_config import (
    ACADEMIC_ONLY_METADATA,
    CONTENT_ROOT,
    LOCAL_OUTPUTS_ERROR,
    ReportConfig,
    load_report_config,
    relative_label,
    targets_local_outputs,
    unknown_route_message,
)
from output_router import FINAL_EXTENSIONS, GLOBAL_OUTPUTS, infer_subject_for_path
from validate_ieee_refs import ValidationResult, validate_ieee
from visual_metadata import validate_visual_manifest

# Visual PDF auditor integration — optional dependency
try:
    # page_issues and default_output_dir come along deliberately: the auditor
    # owns what a finding means and where its artefacts go. Re-deriving either
    # here is how this module used to report the same page as an error and a
    # warning at once, and how a PDF outside the content tree still wrote its
    # renders into it.
    from visual_pdf_auditor import (
        FAILURE,
        WARNING,
        audit_pdf,
        default_output_dir,
        page_issues,
    )
    VISUAL_AUDITOR_AVAILABLE = True
except ImportError:
    VISUAL_AUDITOR_AVAILABLE = False

A4_WIDTH = 595.28
A4_HEIGHT = 841.89

# ---------------------------------------------------------------------------
# Overfull box thresholds
# ---------------------------------------------------------------------------
#
# LaTeX reports every box that does not fit, and a real log carries dozens of
# them. Severity has to be proportional or the signal is buried again.
#
# Below OVERFULL_NOISE_PT the overflow is a sub-point rounding artefact: at
# 300 dpi one point is four pixels, so a fraction of a point is not visible on
# paper and is not worth a line of output.
#
# OVERFULL_CLIPPING_PT is the page margin. Every template compiles with
# `\usepackage[a4paper,margin=2.5cm]{geometry}`, and 2.5 cm is 71.13 pt, so an
# overflow wider than that does not merely intrude into the white margin: it
# reaches the physical paper edge and everything past it is cut off. That is
# lost content, not a typographic blemish, so it is an error. Anything in
# between is visible in the margin but still on the page — a warning.
OVERFULL_NOISE_PT = 1.0
OVERFULL_CLIPPING_PT = 71.13

# Cap on individually listed boxes; the total is always reported.
OVERFULL_REPORT_LIMIT = 8

# LaTeX hard-wraps log lines at `max_print_line` (79 or 80 in practice), so a
# message can be split anywhere. Only a line that long can have been wrapped,
# and only an incomplete message is ever re-joined.
LATEX_LOG_WRAP_MIN_WIDTH = 60
OVERFULL_LOOKAHEAD_LINES = 2

OVERFULL_RE = re.compile(
    r"Overfull \\(?P<kind>[hv])box \((?P<points>\d+(?:\.\d+)?)pt too (?:wide|high)\)"
)
OVERFULL_LINES_RE = re.compile(r"at lines (?P<start>\d+)--(?P<end>\d+)")
OVERFULL_LINE_RE = re.compile(r"at line (?P<start>\d+)\b")
# A page ships out as `[12]`, ` [12<image.png>]` or ` [12`. Version stamps such
# as `[2026/01/01 v1.2]` are excluded by the lookahead.
PAGE_SHIPOUT_RE = re.compile(r"(?:^|[\s)])\[(\d+)(?=[\]\s<{]|$)")


@dataclass(frozen=True)
class OverfullBox:
    """One `Overfull \\hbox/\\vbox` report parsed out of a LaTeX log."""

    kind: str
    points: float
    start_line: int | None = None
    end_line: int | None = None
    page: int | None = None

    def location(self) -> str:
        parts = []
        if self.page is not None:
            parts.append(f"página {self.page}")
        if self.start_line is not None:
            if self.end_line is not None and self.end_line != self.start_line:
                parts.append(f"líneas {self.start_line}--{self.end_line}")
            else:
                parts.append(f"línea {self.start_line}")
        return ", ".join(parts) if parts else "ubicación no indicada en el log"

    def describe(self) -> str:
        return f"\\{self.kind} de {self.points:.2f}pt ({self.location()})"


def _parse_overfull(window: str, page: int | None) -> OverfullBox | None:
    match = OVERFULL_RE.search(window)
    if match is None:
        return None
    tail = window[match.end():]
    lines_match = OVERFULL_LINES_RE.search(tail)
    if lines_match:
        start = int(lines_match.group("start"))
        end = int(lines_match.group("end"))
    else:
        single = OVERFULL_LINE_RE.search(tail)
        start = int(single.group("start")) if single else None
        end = start
    return OverfullBox(
        kind=f"{match.group('kind')}box",
        points=float(match.group("points")),
        start_line=start,
        end_line=end,
        page=page,
    )


def _overfull_is_complete(window: str) -> bool:
    """True when the window holds a whole message, so no unwrapping is needed."""
    match = OVERFULL_RE.search(window)
    if match is None:
        return False
    tail = window[match.end():].rstrip()
    location = OVERFULL_LINES_RE.search(tail) or OVERFULL_LINE_RE.search(tail)
    if location is None:
        # `has occurred while \output is active` carries no source location.
        return "is active" in tail or "has occurred" in tail
    # A location running to the very end of the line may have been cut by the
    # hard wrap (`... at lines 1780--17` continued by `99`), so it counts as
    # complete only when something follows it.
    return location.end() < len(tail)


def parse_overfull_boxes(log: str) -> list[OverfullBox]:
    """Extract every Overfull box report from a LaTeX log.

    Tolerates a missing, empty or truncated log, and re-joins messages that
    LaTeX hard-wrapped mid-sentence. The page is inferred from the shipout
    markers that precede the message; when the log names no page at all (a
    truncated log, or a run that never shipped a page) it is left unknown
    rather than guessed.
    """
    lines = log.splitlines()
    has_shipouts = any(PAGE_SHIPOUT_RE.search(line) for line in lines)
    next_page = 1
    boxes: list[OverfullBox] = []

    for index, line in enumerate(lines):
        if "Overfull" in line:
            window = line
            steps = 0
            while (
                not _overfull_is_complete(window)
                and steps < OVERFULL_LOOKAHEAD_LINES
                and len(lines[index + steps]) >= LATEX_LOG_WRAP_MIN_WIDTH
                and index + steps + 1 < len(lines)
            ):
                steps += 1
                window += lines[index + steps]
            box = _parse_overfull(window, next_page if has_shipouts else None)
            if box is not None:
                boxes.append(box)
        for marker in PAGE_SHIPOUT_RE.findall(line):
            next_page = int(marker) + 1

    return boxes


def overfull_validation(boxes: list[OverfullBox]) -> ValidationResult:
    """Turn parsed Overfull boxes into proportional errors and warnings."""
    result = ValidationResult()
    relevant = [box for box in boxes if box.points >= OVERFULL_NOISE_PT]
    if not relevant:
        return result

    clipping = [box for box in relevant if box.points >= OVERFULL_CLIPPING_PT]
    visible = [box for box in relevant if box.points < OVERFULL_CLIPPING_PT]

    for box in clipping[:OVERFULL_REPORT_LIMIT]:
        result.errors.append(
            f"Contenido cortado en el borde de la hoja: Overfull {box.describe()} "
            f"supera el margen de {OVERFULL_CLIPPING_PT:.0f}pt; el texto que sobra no se imprime"
        )
    if len(clipping) > OVERFULL_REPORT_LIMIT:
        result.errors.append(
            f"Hay {len(clipping)} desbordes que cortan contenido; "
            f"se listaron los primeros {OVERFULL_REPORT_LIMIT}"
        )

    for box in visible[:OVERFULL_REPORT_LIMIT]:
        result.warnings.append(
            f"Overfull {box.describe()} invade el margen; revisar corte de línea, "
            "URL larga o ancho de tabla"
        )
    if len(visible) > OVERFULL_REPORT_LIMIT:
        result.warnings.append(
            f"Hay {len(visible)} desbordes que invaden el margen; "
            f"se listaron los primeros {OVERFULL_REPORT_LIMIT}"
        )
    return result


@dataclass
class ReportValidation:
    errors: list[str] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)
    checks: list[str] = field(default_factory=list)

    def add(self, name: str, result: ValidationResult) -> None:
        self.checks.append(name)
        self.errors.extend(result.errors)
        self.warnings.extend(result.warnings)


def run(cmd: list[str]) -> subprocess.CompletedProcess[str]:
    return subprocess.run(cmd, text=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE, check=False)


def pdf_text_pages(pdf: Path) -> list[str]:
    if not pdf.exists():
        return []
    result = run(["pdftotext", str(pdf), "-"])
    if result.returncode != 0:
        return []
    return result.stdout.split("\f")


def pdfinfo(pdf: Path) -> dict[str, str]:
    if not pdf.exists():
        return {}
    result = run(["pdfinfo", str(pdf)])
    data: dict[str, str] = {}
    for line in result.stdout.splitlines():
        key, sep, value = line.partition(":")
        if sep:
            data[key.strip()] = value.strip()
    return data


def metadata_validation(config: ReportConfig) -> ValidationResult:
    """Check report.yml metadata against the report's document route.

    The required set is route-dependent (see ``ROUTE_REQUIRED_METADATA`` in
    report_config). Routes B–D forbid the academic shell, so demanding
    `subject`/`teacher` there left users with one escape: inventing fake
    academic metadata — exactly what the routing contract exists to prevent.
    """
    result = ValidationResult()
    if not config.route_is_known:
        # No metadata verdict is possible without a route, and guessing one
        # would reintroduce the silent fallback to Route A.
        result.errors.append(unknown_route_message(config.route))
        return result

    meta = config.metadata
    missing = [key for key in config.required_metadata if not str(meta.get(key) or "").strip()]
    if missing:
        result.errors.append("Metadata incompleta en report.yml: " + ", ".join(missing))

    if config.route != "academic":
        declared = [key for key in ACADEMIC_ONLY_METADATA if str(meta.get(key) or "").strip()]
        if declared:
            result.warnings.append(
                f"La ruta '{config.route}' no es académica pero report.yml declara "
                "metadata académica: " + ", ".join(declared)
            )
    return result


def common_validation(config: ReportConfig) -> ValidationResult:
    result = ValidationResult()
    metadata_result = metadata_validation(config)
    result.errors.extend(metadata_result.errors)
    result.warnings.extend(metadata_result.warnings)

    if not config.body_path.exists() and config.backend in {"latex", "html"}:
        result.errors.append(f"No existe body.md: {config.body_path}")

    if config.output_format == "pdf" and not config.pdf_path.exists():
        result.errors.append(f"No existe PDF final esperado: {config.pdf_path}")

    outputs = config.folder / "outputs"
    local_outputs = outputs.resolve()
    global_outputs = GLOBAL_OUTPUTS.resolve()
    final_pdf = config.pdf_path.resolve()
    if config.output_format == "pdf" and config.pdf_path.exists():
        in_local_outputs = local_outputs in final_pdf.parents
        in_global_outputs = global_outputs in final_pdf.parents
        if not (in_local_outputs or in_global_outputs):
            expected_global_exists = False
            if config.publish_global:
                subject_slug = infer_subject_for_path(config.pdf_path, config.metadata)
                if subject_slug:
                    expected_global = GLOBAL_OUTPUTS / subject_slug / config.pdf_path.name
                    expected_global_exists = expected_global.exists()
            if not expected_global_exists:
                result.warnings.append("El PDF final no está dentro de outputs/")
        if targets_local_outputs(config):
            # Defence in depth: load_report_config already refuses a declared
            # path here, so reaching this branch means the PDF landed in the
            # forbidden folder some other way (an implicit default, a stale
            # copy, a builder writing outside its configuration).
            #
            # The condition is the shared predicate, not a second spelling of
            # it: this branch used to repeat the underscore exemption inline,
            # which is two places to keep in step for one rule.
            result.errors.append(LOCAL_OUTPUTS_ERROR)

    if outputs.exists():
        dirty = [p.name for p in outputs.iterdir() if p.is_file() and p.suffix.lower() not in {".pdf", ".docx"}]
        if dirty:
            result.warnings.append("outputs/ contiene intermedios; deberían ir a backups/ o build/: " + ", ".join(dirty[:8]))

    if config.publish_global and config.output_format == "pdf" and config.pdf_path.exists():
        subject_slug = infer_subject_for_path(config.pdf_path, config.metadata)
        if subject_slug:
            expected_global = GLOBAL_OUTPUTS / subject_slug / config.pdf_path.name
            if not expected_global.exists():
                result.warnings.append(
                    "No existe copia global filtrada por materia: "
                    + relative_label(expected_global, GLOBAL_OUTPUTS.parent)
                )
        else:
            result.warnings.append("No pude inferir la materia para publicar en outputs/<materia>/")

    if GLOBAL_OUTPUTS.exists():
        loose_items = [p.name for p in GLOBAL_OUTPUTS.iterdir() if p.is_file() and p.name != ".gitkeep"]
        loose_finals = [name for name in loose_items if Path(name).suffix.lower() in FINAL_EXTENSIONS]
        loose_intermediates = [name for name in loose_items if Path(name).suffix.lower() not in FINAL_EXTENSIONS]
        non_delivery = [
            str(p.relative_to(GLOBAL_OUTPUTS))
            for p in GLOBAL_OUTPUTS.rglob("*")
            if p.is_file() and p.name != ".gitkeep" and p.suffix.lower() not in FINAL_EXTENSIONS
        ]
        if loose_finals:
            result.warnings.append("outputs/ global tiene finales sueltos; deben ir en outputs/<materia>/: " + ", ".join(loose_finals[:8]))
        if loose_intermediates:
            result.warnings.append("outputs/ global tiene intermedios sueltos; deben ir a backups/: " + ", ".join(loose_intermediates[:8]))
        if non_delivery:
            result.warnings.append("outputs/ global contiene archivos que no son PDF/DOCX; deben ir a backups/: " + ", ".join(non_delivery[:8]))

    text = "\n".join(pdf_text_pages(config.pdf_path)) if config.pdf_path.exists() else ""
    banned = config.academic_value("conclusions", "banned_openers", default=["Se concluye", "En conclusión", "En conclusion"])
    found = [item for item in banned if item.lower() in text.lower()]
    if found:
        result.errors.append("Conclusión inicia/usa fórmulas prohibidas: " + ", ".join(found))

    if config.pdf_path.exists():
        pages = pdf_text_pages(config.pdf_path)
        blank_pages = [str(i + 1) for i, page in enumerate(pages) if i < len(pages) - 1 and len(page.strip()) < 20]
        if blank_pages:
            result.errors.append("Posibles páginas vacías accidentales: " + ", ".join(blank_pages))

        if config.academic_value("cover", "required", default=True) and len(pages) < 2:
            result.errors.append("La portada es obligatoria pero el PDF tiene menos de 2 páginas")

        # CRITICAL: detect LaTeX commands rendered as literal text in PDF
        full_text = "\n".join(pages)
        latex_escaped_patterns = [
            (r"\quad", "\\quad"),
            (r"\qquad", "\\qquad"),
            (r"\frac", "\\frac"),
            (r"\boxed", "\\boxed"),
            (r"\sqrt", "\\sqrt"),
            (r"\pi", "\\pi"),
            (r"\sin", "\\sin"),
            (r"\cos", "\\cos"),
            (r"\tan", "\\tan"),
            (r"\pm", "\\pm"),
            (r"\cdot", "\\cdot"),
            (r"\Longrightarrow", "\\Longrightarrow"),
            (r"\left", "\\left"),
            (r"\right", "\\right"),
            (r"\begin", "\\begin"),
            (r"\end", "\\end"),
            (r"\textbackslash", "\\textbackslash"),
        ]
        found = []
        for display_name, search_pattern in latex_escaped_patterns:
            count = full_text.count(search_pattern)
            if count > 0:
                found.append(f"{display_name} ({count} vez/veces)")
        if found:
            result.errors.append(
                "COMANDOS LATEX RENDERIZADOS COMO TEXTO LITERAL EN EL PDF: "
                + ", ".join(found) + ". "
                "Indica que latex_escape() en build_latex_report.py está "
                "escapando comandos que deberían estar dentro de $...$ o $$...$$."
            )

    return result


def asset_validation(_config: ReportConfig) -> ValidationResult:
    """Validate that system-level build assets (logo, background) exist.

    This checks the automation root's assets/ folder, not the report folder.
    Every LaTeX pipeline run needs these files; warn early if they drift.
    """
    result = ValidationResult()
    from build_latex_report import ASSETS_DIR, EXPECTED_ASSETS

    for filename, description in EXPECTED_ASSETS:
        path = ASSETS_DIR / filename
        if not path.exists():
            result.errors.append(
                f"Asset faltante: {description} ({filename}) — esperado en {ASSETS_DIR}"
            )
        elif path.stat().st_size < 1_000:
            result.errors.append(f"Asset sospechosamente pequeño: {path}")
    return result


def pdf_layout_validation(config: ReportConfig) -> ValidationResult:
    result = ValidationResult()
    if not config.pdf_path.exists():
        result.errors.append(f"No existe PDF para validar layout: {config.pdf_path}")
        return result

    info = pdfinfo(config.pdf_path)
    page_count = int(info.get("Pages", "0") or 0)
    if page_count < 1:
        result.errors.append("PDF sin páginas o pdfinfo no pudo leer conteo")

    page_size = info.get("Page size", "")
    numbers = [float(n) for n in re.findall(r"\d+(?:\.\d+)?", page_size)[:2]]
    if len(numbers) == 2:
        w, h = numbers
        is_a4 = abs(w - A4_WIDTH) < 3 and abs(h - A4_HEIGHT) < 3
        is_a4_landscape = abs(w - A4_HEIGHT) < 3 and abs(h - A4_WIDTH) < 3
        if not (is_a4 or (config.backend == "visual" and is_a4_landscape)):
            result.warnings.append(f"Tamaño de página no parece A4 estándar: {page_size}")

    try:
        images = run(["pdfimages", "-list", str(config.pdf_path)]).stdout.splitlines()
        image_rows = [line for line in images if re.match(r"^\s*\d+\s+\d+\s+image", line)]
        if config.academic_value("cover", "logo_required", default=True) and not image_rows:
            result.errors.append("PDF no parece tener imágenes embebidas; revisar logo UNL")
    except FileNotFoundError:
        result.warnings.append("pdfimages no disponible; no se validó logo/imágenes")

    pages = pdf_text_pages(config.pdf_path)
    body_page = config.academic_value("cover", "body_starts_on_page", default=2)
    body_page_index = body_page - 1  # 0-indexed
    # The cover/body boundary only exists when there IS a cover. A report that
    # declared `cover: {required: false}` starts its body on page 1, and the
    # markers below are academic Spanish — a business report opening on
    # "Resumen ejecutivo" can never match one, so the warning fired on every
    # build and taught the reader to skip warnings entirely.
    has_cover = bool(config.academic_value("cover", "required", default=True))
    if has_cover and len(pages) > body_page_index:
        first = pages[0].lower()
        body_content = pages[body_page_index].lower()
        body_markers = [
            "introducción",
            "introduccion",
            "tema",
            "antecedentes",
            "descripción",
            "descripcion",
            "desarrollo",
            "ejercicio",
            "ejercicios",
        ]
        body_marker_pattern = r"\b(" + "|".join(re.escape(marker) for marker in body_markers) + r")\b"
        cover_body_markers = [
            "introducción",
            "introduccion",
            "tema",
            "antecedentes",
            "descripción",
            "descripcion",
            "desarrollo",
        ]
        cover_body_marker_pattern = r"\b(" + "|".join(re.escape(marker) for marker in cover_body_markers) + r")\b"
        if re.search(cover_body_marker_pattern, first) and not config.raw.get("allow_body_on_cover", False):
            result.errors.append("La portada parece mezclada con el cuerpo; el cuerpo debe iniciar en página 2")
        if not re.search(body_marker_pattern, body_content) and config.backend == "latex":
            result.warnings.append(f"No detecté inicio claro del cuerpo en página {body_page}; revisar portada/cuerpo")

    # Orphan headings are a pagination defect, not a cover concern: this loop
    # only ever needed `pages`. It used to sit inside the cover/body branch,
    # which meant a report without a cover — or one whose body legitimately
    # starts later — silently lost the check.
    for idx, page in enumerate(pages[:-1], start=1):
        lines = [line.strip() for line in page.splitlines() if line.strip()]
        if not lines:
            continue
        last = lines[-1]
        looks_heading = (
            bool(re.match(r"^(\d+(?:\.\d+)*)?\s*[A-ZÁÉÍÓÚÑ][\wÁÉÍÓÚÑáéíóúñ\s]{3,70}$", last))
            and not last.endswith((".", ":", ";", ","))
        )
        if looks_heading:
            # Reduce false positives from pdftotext fragments: a heading
            # should have a number prefix, span multiple words, or be at
            # least 8 characters (single short words like "Manten" are
            # typically split-word artifacts, not real headings).
            has_number = bool(re.match(r"^\d+(?:\.\d+)*", last))
            has_space = " " in last.strip()
            if not (has_number or has_space or len(last.strip()) >= 8):
                looks_heading = False
        if looks_heading:
            result.errors.append(f"Posible título huérfano al final de página {idx}: {last}")

    return result


def latex_log_validation(config: ReportConfig) -> ValidationResult:
    result = ValidationResult()
    log = config.log_path.read_text(encoding="utf-8", errors="ignore") if config.log_path.exists() else ""
    tex = config.tex_path.read_text(encoding="utf-8", errors="ignore") if config.tex_path.exists() else ""
    if not config.tex_path.exists():
        result.errors.append(f"No existe main.tex: {config.tex_path}")
        return result

    required_snippets = ["hyperref", "Needspace", "onehalfspacing", "parindent"]
    missing = [snippet for snippet in required_snippets if snippet not in tex]
    if missing:
        result.errors.append("LaTeX template sin reglas obligatorias: " + ", ".join(missing))

    if log:
        fatal_patterns = [
            r"! LaTeX Error",
            r"Undefined control sequence",
            r"Citation `[^']+' undefined",
            r"Reference `[^']+' undefined",
            r"There were undefined (?:references|citations)",
            r"Package biblatex Warning: Please \(re\)run Biber",
        ]
        found = [pattern for pattern in fatal_patterns if re.search(pattern, log)]
        if found:
            result.errors.append("Log LaTeX contiene errores/referencias sin resolver")
        overfull = overfull_validation(parse_overfull_boxes(log))
        result.errors.extend(overfull.errors)
        result.warnings.extend(overfull.warnings)
    else:
        result.warnings.append("No existe log LaTeX todavía; se validó solo el .tex")

    return result


def source_layout_validation(config: ReportConfig) -> ValidationResult:
    result = ValidationResult()
    body = config.body_path.read_text(encoding="utf-8", errors="ignore") if config.body_path.exists() else ""
    bad_bold_titles = [line for line in body.splitlines() if re.match(r"^\s*\*\*[^*]{4,}\*\*\s*$", line)]
    if bad_bold_titles:
        result.errors.append("Usar headings Markdown (#, ##) para títulos, no negrita manual: " + bad_bold_titles[0][:80])
    return result


def visual_validation(config: ReportConfig) -> ValidationResult:
    result = ValidationResult()
    figures_yml = config.folder / "figures" / "figures.yml"
    if not figures_yml.exists():
        figures_yml = config.folder / "figures.yml"
    if not figures_yml.exists():
        result.warnings.append("Trabajo visual sin figures.yml; no se pudo validar captions/fuentes de figuras")
        return result
    metadata_result = validate_visual_manifest(figures_yml.parent, figures_yml)
    result.errors.extend(metadata_result.errors)
    result.warnings.extend(metadata_result.warnings)
    for figure_path in figures_yml.parent.rglob("*"):
        if figure_path.is_file() and figure_path.suffix.lower() in {".svg", ".png", ".pdf"} and figure_path.stat().st_size < 4_000:
            result.warnings.append(f"Figura parece demasiado pequeña: {figure_path}")
    return result


def docx_validation(config: ReportConfig) -> ValidationResult:
    """Validate the DOCX artefact itself, not merely that a file exists.

    Severity follows the enforcement key documented at the top of
    ``templates/academic_format.yml``, so this validator does not invent a
    stricter contract than the format file declares:

    * lost or unreadable content, a degenerate table and a figure that never
      made it into the package are ERRORS — they are missing deliverable, not
      styling;
    * ``page.size`` is annotated ``[ENFORCED]``, so a document that is not the
      declared size is an ERROR;
    * margins, body font and body size are annotated ``[GUIDANCE]``/
      ``[TEST-GATED]`` ("set in the .tex template"), so a mismatch is a
      WARNING. A DOCX restyled by a task-specific script (restyle_docx_aa1.py)
      legitimately differs there and must not be failed for it.
    """
    result = ValidationResult()
    if not config.docx_path.exists():
        result.errors.append(f"No existe DOCX final esperado: {config.docx_path}")
        return result
    if config.docx_path.stat().st_size < 10_000:
        result.warnings.append("DOCX demasiado pequeño; revisar que no esté vacío")

    try:
        from docx import Document as DocxDocument
    except ImportError:
        result.warnings.append(
            "python-docx no está instalado; no pude validar el contenido del DOCX"
        )
        return result

    try:
        document = DocxDocument(str(config.docx_path))
    except Exception as exc:
        result.errors.append(
            f"No se pudo abrir el DOCX ({config.docx_path}): {exc}. "
            "El archivo está corrupto o no es un documento de Word."
        )
        return result

    filled = [p for p in document.paragraphs if p.text.strip()]
    if not filled and not document.tables:
        result.errors.append(
            f"El DOCX está sin contenido: {config.docx_path} no tiene párrafos "
            "con texto ni tablas"
        )
        return result

    if not any(p.style is not None and p.style.name.startswith("Heading") for p in document.paragraphs):
        result.warnings.append(
            "El DOCX no tiene títulos con estilo Heading; el documento no será navegable en Word"
        )

    section = document.sections[0]
    expected_size = str(config.academic_value("page", "size", default="A4")).strip().upper()
    width_cm = section.page_width / 360000
    height_cm = section.page_height / 360000
    if expected_size == "A4" and not (
        abs(width_cm - 21.0) < 0.2 and abs(height_cm - 29.7) < 0.2
    ):
        result.errors.append(
            f"El DOCX no está en A4: {width_cm:.2f} x {height_cm:.2f} cm "
            f"(esperado 21.00 x 29.70 cm)"
        )

    allowed_margins = config.academic_value("page", "margins", "allowed_cm", default=[2.5, 3.0])
    margins_cm = {
        "superior": section.top_margin / 360000,
        "inferior": section.bottom_margin / 360000,
        "izquierdo": section.left_margin / 360000,
        "derecho": section.right_margin / 360000,
    }
    off_contract = [
        f"{name} {value:.2f} cm"
        for name, value in margins_cm.items()
        if not any(abs(value - float(allowed)) < 0.1 for allowed in allowed_margins)
    ]
    if off_contract:
        result.warnings.append(
            "Márgenes del DOCX fuera de los permitidos "
            f"({', '.join(str(m) for m in allowed_margins)} cm): " + ", ".join(off_contract)
        )

    normal = document.styles["Normal"]
    preference = config.academic_value("body_text", "font_preference", default=[]) or []
    expected_font = str(preference[0]) if preference else "Times New Roman"
    if normal.font.name and normal.font.name != expected_font:
        result.warnings.append(
            f"Fuente base del DOCX: '{normal.font.name}'; el formato declara '{expected_font}'"
        )
    expected_pt = float(config.academic_value("body_text", "size_pt", default=12))
    if normal.font.size is not None and abs(normal.font.size.pt - expected_pt) > 0.1:
        result.warnings.append(
            f"Tamaño base del DOCX: {normal.font.size.pt} pt; el formato declara {expected_pt} pt"
        )

    for index, table in enumerate(document.tables, start=1):
        if not table.rows or not table.columns:
            result.errors.append(f"Tabla {index} del DOCX quedó sin filas o sin columnas")
            continue
        empty_rows = [
            position
            for position, row in enumerate(table.rows, start=1)
            if not any(cell.text.strip() for cell in row.cells)
        ]
        if empty_rows:
            result.errors.append(
                f"Tabla {index} del DOCX tiene filas totalmente vacías: "
                + ", ".join(str(position) for position in empty_rows)
            )

    if config.body_path.exists():
        body = config.body_path.read_text(encoding="utf-8", errors="ignore")
        local_figures = [
            reference
            for reference in re.findall(r"!\[[^\]]*\]\(([^)]+)\)", body)
            if not reference.strip().startswith(("http://", "https://", "data:"))
        ]
        if local_figures and not document.inline_shapes:
            result.errors.append(
                "El body.md referencia imágenes pero el DOCX no tiene ninguna incrustada: "
                + ", ".join(local_figures[:5])
            )

    return result


def visual_pdf_validation(config: ReportConfig) -> ValidationResult:
    """Validate PDF visual quality using visual_pdf_auditor.py.
    
    Runs the heuristic page-image analysis (pdftoppm + PIL) on the final PDF.
    Requires poppler-utils (pdftoppm) and Pillow.
    Only runs when config.output_format == 'pdf' and the PDF exists.
    """
    result = ValidationResult()

    if config.output_format != "pdf":
        result.warnings.append("visual_pdf: saltado — formato no es PDF")
        return result

    pdf = config.pdf_path
    if not pdf.exists():
        result.errors.append(f"visual_pdf: no existe PDF para auditar: {pdf}")
        return result

    if not VISUAL_AUDITOR_AVAILABLE:
        result.warnings.append(
            "visual_pdf: visual_pdf_auditor no disponible; "
            "ejecutá manual: python tools/visual_pdf_auditor.py <pdf>"
        )
        return result

    try:
        audit_result = audit_pdf(pdf, default_output_dir(pdf))

        for finding in audit_result.findings:
            for issue in page_issues(finding, audit_result.total_pages):
                line = f"visual_pdf [Página {finding.page}]: {issue.tag}"
                if issue.detail:
                    line += f" ({issue.detail})"
                if issue.level == FAILURE:
                    result.errors.append(line)
                elif issue.level == WARNING:
                    result.warnings.append(line)
                # INFO stays out: it is context for a human reading the audit,
                # not something a build gate should speak about.

    except Exception as exc:
        result.warnings.append(f"visual_pdf: error ejecutando auditor: {exc}")

    return result


def write_quality_report(config: ReportConfig, validation: ReportValidation) -> None:
    path = config.quality_report_path
    path.parent.mkdir(parents=True, exist_ok=True)
    status = "APROBADO" if not validation.errors else "FALLÓ"
    lines = [
        f"# Quality report — {config.folder.name}",
        "",
        "## Resultado",
        status,
        "",
        "## Validaciones ejecutadas",
    ]
    lines.extend(f"- {check}" for check in validation.checks)
    if validation.errors:
        lines.extend(["", "## Errores", *[f"- {error}" for error in validation.errors]])
    if validation.warnings:
        lines.extend(["", "## Advertencias", *[f"- {warning}" for warning in validation.warnings]])
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def validate(config: ReportConfig) -> ReportValidation:
    validators = config.validators
    validation = ReportValidation()
    if validators.get("common", True):
        validation.add("common", common_validation(config))
    if (
        validators.get("common", True)
        and config.backend == "latex"
    ):
        validation.add("assets", asset_validation(config))
    if validators.get("pdf_layout", False):
        validation.add("pdf_layout", pdf_layout_validation(config))
    if validators.get("ieee", True):
        validation.add("ieee", validate_ieee(config))
    if validators.get("latex", False):
        validation.add("source_layout", source_layout_validation(config))
    if validators.get("latex_log", False):
        validation.add("latex_log", latex_log_validation(config))
    if validators.get("visual", False):
        validation.add("visual", visual_validation(config))
    if validators.get("visual_pdf", False):
        validation.add("visual_pdf", visual_pdf_validation(config))
    if validators.get("docx", False):
        validation.add("docx", docx_validation(config))
    write_quality_report(config, validation)
    return validation


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("folder", type=Path, help="Carpeta del reporte con report.yml")
    args = parser.parse_args()
    config = load_report_config(args.folder)
    result = validate(config)
    if result.errors:
        raise SystemExit("VALIDATION FAILED:\n- " + "\n- ".join(result.errors) + f"\n\nReporte: {config.quality_report_path}")
    if result.warnings:
        print("VALIDATION PASSED WITH WARNINGS:\n- " + "\n- ".join(result.warnings))
    else:
        print("VALIDATION PASSED")
    print(f"Reporte: {config.quality_report_path}")


if __name__ == "__main__":
    main()
