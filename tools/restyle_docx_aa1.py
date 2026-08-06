#!/usr/bin/env python3
"""Restyle an academic DOCX to match Alejandro's 10/10 AA1 visual profile.

Usage:
  python3 tools/restyle_docx_aa1.py input.docx --output outputs/fixed.docx --logo assets/unl-logo-aa1.png

Input expectations:
- first table: cover metadata with label/value rows
- body starts at the first numbered heading paragraph
- optional second table: comparative/technical body table
"""
from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

FONT = "Times New Roman"


def set_run(run, size=None, bold=None, italic=None, color=None, font=FONT):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill="F2F2F2"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, val="single", sz="8", color="000000"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), val)
        node.set(qn("w:sz"), sz)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def clear_cell_borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "nil")


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, pct=10000):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(pct))
    tbl_w.set(qn("w:type"), "pct")


def setup_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(12)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.first_line_indent = Cm(1.27)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.12

    for style_name, size in [("Heading 1", 14), ("Heading 2", 12)]:
        st = doc.styles[style_name]
        st.font.name = FONT
        st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor(0, 0, 0)
        st.paragraph_format.space_before = Pt(12)
        st.paragraph_format.space_after = Pt(6)
        st.paragraph_format.line_spacing = 1.0


def add_center(doc, text="", size=12, bold=False, italic=False, after=3, before=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    set_run(r, size=size, bold=bold, italic=italic)
    return p


def add_header(section, logo_path: Path):
    header = section.header
    for p in header.paragraphs:
        p._element.getparent().remove(p._element)
    table = header.add_table(rows=1, cols=3, width=Inches(6.6))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Inches(1.55), Inches(4.1), Inches(1.15)]
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = widths[idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            clear_cell_borders(cell)
            set_cell_margins(cell, top=0, bottom=0, start=30, end=30)

    p = table.cell(0, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run().add_picture(str(logo_path), width=Inches(1.28))

    p = table.cell(0, 1).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, text in enumerate([
        "UNIVERSIDAD NACIONAL DE LOJA",
        "FACULTAD DE ENERGÍA, LAS INDUSTRIAS Y LOS RECURSOS",
        "NATURALES NO RENOVABLES",
        "CARRERA DE COMPUTACIÓN",
    ]):
        if i:
            p.add_run().add_break()
        set_run(p.add_run(text), size=8.2, bold=True)

    right = table.cell(0, 2)
    set_cell_borders(right, sz="8")
    p = right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, text in enumerate(["CARRERA DE", "COMPUTACIÓN", "HE-CIS-2022"]):
        if i:
            p.add_run().add_break()
        set_run(p.add_run(text), size=8.7, bold=True)

    line = header.add_paragraph()
    line.paragraph_format.space_before = Pt(1)
    line.paragraph_format.space_after = Pt(0)
    p_pr = line._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "18")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_footer(section):
    footer = section.footer
    for p in footer.paragraphs:
        p._element.getparent().remove(p._element)
    p = footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    set_run(p.add_run("Educamos para "), size=9.5, italic=True)
    set_run(p.add_run("Transformar"), size=9.5, italic=True, bold=True)


def add_cover(doc, meta, logo_path: Path):
    add_center(doc, "", after=8)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(logo_path), width=Inches(1.65))
    p.paragraph_format.space_after = Pt(18)

    add_center(doc, "UNIVERSIDAD NACIONAL DE LOJA", size=15, bold=True, after=4)
    add_center(doc, "FACULTAD DE LA ENERGÍA, LAS INDUSTRIAS Y", size=12.5, bold=True, after=0)
    add_center(doc, "LOS RECURSOS NATURALES NO RENOVABLES", size=12.5, bold=True, after=0)
    add_center(doc, "CARRERA COMPUTACIÓN", size=12.5, bold=True, after=22)

    kind = meta.get("Tipo", "Ensayo")
    title_top = "Ensayo" if kind.lower().startswith("ensayo") else "Aprendizaje Autónomo"
    add_center(doc, title_top, size=20, bold=True, after=10)
    if title_top == "Ensayo":
        add_center(doc, meta.get("Título", "").upper(), size=12, bold=True, after=18)
    else:
        add_center(doc, "Actividad Nro. 1", size=17, bold=True, after=18)

    rows = [
        ("Asignatura", meta.get("Asignatura", "")),
        ("Título", meta.get("Título", "")),
        ("Tipo", meta.get("Tipo", "")),
        ("Docente", meta.get("Docente", "")),
        ("Estudiante", meta.get("Estudiante", "")),
        ("Paralelo", meta.get("Paralelo", "")),
        ("Período\nAcadémico", meta.get("Período Académico", "")),
        ("Fecha", meta.get("Fecha", "")),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, (label, value) in enumerate(rows):
        c0, c1 = table.rows[i].cells
        c0.width = Inches(1.8)
        c1.width = Inches(4.2)
        for cell in (c0, c1):
            set_cell_borders(cell, sz="6")
            set_cell_margins(cell, top=95, bottom=95, start=130, end=130)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(c0, "F2F2F2")
        p0 = c0.paragraphs[0]
        for part_i, part in enumerate(label.split("\n")):
            if part_i:
                p0.add_run().add_break()
            set_run(p0.add_run(part), size=12, bold=True)
        set_run(c1.paragraphs[0].add_run(value), size=12)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(34)
    add_center(doc, 'Ciudad Universitaria "Guillermo Falconí Espinosa"', size=10.5, italic=True, after=0)
    add_center(doc, "Universidad Nacional de Loja", size=10.5, italic=True, after=0)
    doc.add_page_break()


def format_body_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_width(table, 10000)
    for row_i, row in enumerate(table.rows):
        for cell in row.cells:
            set_cell_borders(cell, sz="6")
            set_cell_margins(cell, top=70, bottom=70, start=85, end=85)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            if row_i == 0:
                set_cell_shading(cell, "EDEDED")
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_after = Pt(0)
                for r in p.runs:
                    set_run(r, size=9.2, bold=True if row_i == 0 else None)


def copy_paragraph(doc, src_p):
    text = src_p.text
    style = src_p.style.name
    if style == "Heading 1":
        p = doc.add_paragraph(style="Heading 1")
        set_run(p.add_run(text), size=14, bold=True, color="000000")
        return p
    if style == "Heading 2":
        p = doc.add_paragraph(style="Heading 2")
        set_run(p.add_run(text), size=12, bold=True, color="000000")
        return p

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.12
    if text.startswith("Tabla "):
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(4)
        set_run(p.add_run(text), size=10, italic=True)
        return p
    size = 10.5 if text.startswith("[") else 12
    if text.startswith("["):
        p.paragraph_format.left_indent = Cm(0.75)
        p.paragraph_format.first_line_indent = Cm(-0.75)
    for src_r in src_p.runs:
        set_run(p.add_run(src_r.text), size=size, bold=src_r.bold, italic=src_r.italic)
    if not src_p.runs and text:
        set_run(p.add_run(text), size=size)
    return p


def find_body_start(paragraphs):
    for i, p in enumerate(paragraphs):
        if p.style.name.startswith("Heading") and p.text.strip().startswith("1."):
            return i
    return 0


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("input", type=Path)
    parser.add_argument("--output", type=Path, required=True)
    parser.add_argument("--logo", type=Path, default=Path("assets/unl-logo-aa1.png"))
    args = parser.parse_args()

    src = Document(args.input)
    if not src.tables:
        raise SystemExit("Expected first metadata table in the input DOCX.")
    meta = {}
    for row in src.tables[0].rows:
        if len(row.cells) >= 2:
            meta[row.cells[0].text.strip()] = row.cells[1].text.strip()

    doc = Document()
    setup_styles(doc)
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)
    sec.header_distance = Cm(0.8)
    sec.footer_distance = Cm(0.8)
    sec.different_first_page_header_footer = True
    add_header(sec, args.logo)
    add_footer(sec)
    add_cover(doc, meta, args.logo)

    body_start = find_body_start(src.paragraphs)
    inserted_table = False
    for p in src.paragraphs[body_start:]:
        text = p.text.strip()
        if not text:
            continue
        copy_paragraph(doc, p)
        if text.startswith("Tabla ") and len(src.tables) > 1 and not inserted_table:
            new_table = doc.add_table(rows=0, cols=len(src.tables[1].columns))
            for row in src.tables[1].rows:
                cells = new_table.add_row().cells
                for i, cell in enumerate(row.cells):
                    cells[i].text = cell.text
            format_body_table(new_table)
            inserted_table = True
            doc.add_paragraph()

    args.output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(args.output)
    print(f"DOCX generado: {args.output}")


if __name__ == "__main__":
    main()
